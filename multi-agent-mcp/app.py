
from flask import Flask, request, jsonify, send_from_directory, send_file, render_template, Response, session, redirect, make_response
from flask_cors import CORS
import time
import sys
import os
import logging
import io
import threading
import uuid
import base64
import html
import json
import re
from datetime import datetime
from typing import Optional
import boto3
from botocore.exceptions import ClientError

# Environment / secrets (order matters):
# 1) .env from project directory — primary for local dev and docker --env-file
# 2) config_secure — optional embedded defaults for frozen builds only; must not override .env
# 3) AWS Secrets Manager — optional overlay when AWS_SECRETS_MANAGER_SECRET_ID is set
try:
    from dotenv import load_dotenv

    _APP_ROOT = os.path.dirname(os.path.abspath(__file__))
    load_dotenv(os.path.join(_APP_ROOT, ".env"))
except ImportError:
    pass
try:
    from config_secure import load_secure_env

    load_secure_env()
    print("✅ config_secure: filled missing env keys only (.env / shell take precedence)")
except ImportError:
    print("ℹ️  config_secure not present — using .env / process environment only")
    pass
try:
    from tools.aws_secrets_env import load_aws_secrets_manager_into_environ
except ImportError:
    pass
else:
    try:
        load_aws_secrets_manager_into_environ()
    except Exception as e:
        if (os.getenv("AWS_SECRETS_MANAGER_REQUIRED") or "").strip().lower() in (
            "1",
            "true",
            "yes",
            "on",
        ):
            raise
        print(f"⚠️  AWS Secrets Manager (optional): {e}")
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not installed. Download feature will be disabled.")

# Integrated tools
from tools.bedrock_tool import ask_bedrock
#from tools.gemini_tool import ask_gemini
#from tools.llama_tool import ask_llama
from tools.confluence_tool import confluence_search
#from tools.tickets_tool import read_tickets
from tools.history_tool import add_to_history, get_history
#from tools.suggestions_tool import AI_suggestions
from tools.metrics_persistence import DB_PATH

# Import ask_arlochat (GocBedrock) - auto-detects best mode: SDK async or HTTP fallback
try:
    from tools.ask_arlochat import ask_arlo, MCP_SDK_AVAILABLE
    ARLOCHAT_AVAILABLE = True
    if MCP_SDK_AVAILABLE:
        print("✅ GocBedrock MCP loaded (SDK Async mode - Python 3.10+)")
    else:
        print("✅ GocBedrock MCP loaded (HTTP Fallback mode - Python 3.9+)")
except ImportError as e:
    print(f"⚠️  WARNING: GocBedrock import failed: {e}")
    ARLOCHAT_AVAILABLE = False
    MCP_SDK_AVAILABLE = False
    
    # Create a placeholder function
    def ask_arlo(question: str = "") -> str:
        return f"""
        <div style='background-color: #fee; padding: 12px; border-left: 4px solid #f56565; border-radius: 4px; margin: 8px 0;'>
            <p style='margin: 0; color: #c53030;'>
                ❌ <strong>GocBedrock module failed to load</strong><br><br>
                Error: {html.escape(str(e))}<br><br>
                Please check the logs for more details.
            </p>
        </div>
        """

from tools.service_owners import service_owners_search
from tools.noc_kt import noc_kt_search
from tools.read_arlo_status import read_arlo_status
from tools.oncall_support import confluence_oncall_today
from tools.read_versions import read_versions
from tools.deployed_fw_versions import read_deployed_fw_versions
from tools.deployments_calendar import get_grm_deployments
from tools.datadog_dashboards import read_datadog_dashboards, read_datadog_errors_only, read_datadog_adt, read_datadog_adt_errors_only, read_datadog_samsung, read_datadog_samsung_errors_only, read_datadog_redmetrics_us, read_datadog_all_errors, read_datadog_failed_pods, read_datadog_403_errors, search_datadog_dashboards, search_datadog_services
from tools.partner_monitor_tools import (
    read_datadog_cat,
    read_datadog_cat_errors_only,
    read_datadog_comcast,
    read_datadog_comcast_errors_only,
)
from tools.splunk_tool import (
    read_splunk_p0_dashboard,
    read_splunk_p0_cvr_dashboard,
    read_splunk_p0_adt_dashboard,
    read_splunk_p0_us_infra_dashboard,
)
from tools.service_query import extract_service_name_from_query
from tools.pagerduty_tool import get_pagerduty_incidents
from tools.pagerduty_analytics import get_pagerduty_analytics
from tools.pagerduty_insights import get_pagerduty_insights
from tools.grafana_dashboards import get_grafana_dns_mapper, get_grafana_savant_z2, get_grafana_dashboard_list
from tools.sentinel_certificates import read_sentinel_certificates
from tools.piranha_employees import piranha_employee_lookup
from tools.slack_http import format_slack_connection_error, post_incoming_webhook, post_slack_api

# 📋 Logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler("agent_tool_logs.log", encoding="utf-8"),
        logging.StreamHandler(sys.stdout)
    ]
)

# ✅ Tools
TOOLS = {
    #"Wiki": {"description": "Read workarounds from Confluece", "function": read_tickets},
    "Wiki": {"description": "Read documents from Arlo confluence", "function": confluence_search},
    "Owners": {"description": "Verify who owns each service", "function": service_owners_search},
    "Arlo_Versions": {"description": "Read version information from versions.arlocloud.com", "function": read_versions},
    "Deployed_FW_Versions": {"description": "Read deployed firmware/version matrix from deployed-fw-versions.arlocloud.com", "function": read_deployed_fw_versions},
    "Sentinel_SSL": {"description": "Monitor SSL/TLS certificates from sentinel.arlocloud.com — expired and expiring soon", "function": read_sentinel_certificates},
    "Piranha_Employees": {"description": "Look up employee team, title, and manager from Piranha EngiHub (piranha.arlo.com)", "function": piranha_employee_lookup},
    "DD_Search": {"description": "Search and list Datadog dashboards by name/query", "function": search_datadog_dashboards},
    "DD_Services": {"description": "Search Datadog APM services (backend-*, api-*, etc.)", "function": search_datadog_services},
    "DD_Red_Metrics": {"description": "List and search Datadog dashboards", "function": read_datadog_dashboards},
    "DD_Red_ADT": {"description": "Show RED Metrics - ADT dashboard from Datadog", "function": read_datadog_adt},
    "DD_Red_Samsung": {"description": "Show RED Metrics - Samsung network dashboard from Datadog", "function": read_datadog_samsung},
    "DD_Red_CAT": {"description": "Show RED Metrics - CAT partner network dashboard from Datadog", "function": read_datadog_cat},
    "DD_Red_Comcast": {"description": "Show RED Metrics - Comcast partner network dashboard from Datadog", "function": read_datadog_comcast},
    "DD_Red_Metrics_US": {"description": "Show RED Metrics - US region dashboard from Datadog", "function": read_datadog_redmetrics_us},
    "DD_Errors": {"description": "Show services with errors > 0 from RED Metrics & ADT dashboards", "function": read_datadog_all_errors},
    "DD_Samsung_Errors": {"description": "Show Samsung network services with errors > 0", "function": read_datadog_samsung_errors_only},
    "DD_CAT_Errors": {"description": "Show CAT partner network services with errors > 0", "function": read_datadog_cat_errors_only},
    "DD_Comcast_Errors": {"description": "Show Comcast partner network services with errors > 0", "function": read_datadog_comcast_errors_only},
    "DD_Failed_Pods": {"description": "Monitor Kubernetes pods with failures (ImagePullBackOff, CrashLoop) causing 4xx/5xx errors", "function": read_datadog_failed_pods},
    "DD_403_Errors": {"description": "Monitor 403 Forbidden errors from APM traces (Artifactory, authentication issues)", "function": read_datadog_403_errors},
    "P0_Streaming": {"description": "Show P0 Streaming dashboard from Splunk", "function": read_splunk_p0_dashboard},
    "P0_CVR_Streaming": {"description": "Show P0 CVR Streaming dashboard from Splunk", "function": read_splunk_p0_cvr_dashboard},
    "P0_ADT_Streaming": {"description": "Show P0 ADT Streaming dashboard from Splunk", "function": read_splunk_p0_adt_dashboard},
    "P0_Streaming_US": {"description": "Show P0 Streaming US infra dashboard from Splunk", "function": read_splunk_p0_us_infra_dashboard},
    "Grafana_DNS_Mapper": {"description": "Monitor DNS Mapper IP usage for HMS/CVR streaming (z4)", "function": get_grafana_dns_mapper},
    "Grafana_Savant_z2": {"description": "Monitor Savant infrastructure in Harlem datacenter (z2)", "function": get_grafana_savant_z2},
    "Holiday_Oncall": {"description": "Get on-call schedule for holidays", "function": confluence_oncall_today},
    "PagerDuty": {"description": "Get active incidents from PagerDuty", "function": get_pagerduty_incidents},
    "PagerDuty_Dashboards": {"description": "Show PagerDuty analytics with charts and metrics", "function": get_pagerduty_analytics},
    "PagerDuty_Insights": {"description": "Show incident activity insights and trends", "function": get_pagerduty_insights},
    "Ask_Bedrock": {"description": "Ask AWS Bedrock (Claude 3.5 Sonnet) for AI-powered responses", "function": ask_bedrock},
    "Bedrock_Report": {"description": "AI-powered comprehensive analysis and synthesis", "function": ask_arlo},
}
registered_tools = [(name, tool["description"]) for name, tool in TOOLS.items()]

# Tools that always receive the full user query (MCP / synthesis needs full intent).
_TOOLS_KEEP_FULL_QUERY = frozenset({"Ask_Bedrock", "Bedrock_Report"})
# Tools that ignore service filter and show org-wide data.
_TOOLS_GLOBAL_SCOPE = frozenset({"PagerDuty_Dashboards", "PagerDuty_Insights"})
_TOOLS_WITH_TIMERANGE = frozenset(
    {
        "DD_Search",
        "DD_Services",
        "DD_Red_Metrics",
        "DD_Errors",
        "DD_Red_ADT",
        "DD_Red_Samsung",
        "DD_Red_CAT",
        "DD_Red_Comcast",
        "DD_Red_Metrics_US",
        "DD_Samsung_Errors",
        "DD_CAT_Errors",
        "DD_Comcast_Errors",
        "DD_Failed_Pods",
        "DD_403_Errors",
        "P0_Streaming",
        "P0_CVR_Streaming",
        "P0_ADT_Streaming",
        "P0_Streaming_US",
        "Grafana_DNS_Mapper",
        "Grafana_Savant_z2",
    }
)

_bedrock_runtime_lock = threading.Lock()
_bedrock_runtime_client = None


def _get_bedrock_runtime_client():
    """Reuse boto3 client across requests (thread-safe for typical invoke_model usage)."""
    global _bedrock_runtime_client
    if _bedrock_runtime_client is not None:
        return _bedrock_runtime_client
    bedrock_api_key = os.getenv("BEDROCK_API_KEY")
    if not bedrock_api_key:
        return None
    with _bedrock_runtime_lock:
        if _bedrock_runtime_client is not None:
            return _bedrock_runtime_client
        _bedrock_runtime_client = boto3.client(
            service_name="bedrock-runtime",
            region_name="us-east-1",
            aws_access_key_id=bedrock_api_key.split(":")[0] if ":" in bedrock_api_key else bedrock_api_key,
            aws_secret_access_key=bedrock_api_key.split(":")[1]
            if ":" in bedrock_api_key and len(bedrock_api_key.split(":")) > 1
            else "",
        )
        return _bedrock_runtime_client


def _monitoring_tool_input_from_analysis(user_query: str, analysis: dict) -> str:
    if analysis.get("is_general_query"):
        return ""
    svc = (analysis.get("service_name") or "").strip()
    if svc:
        return svc
    return extract_service_name_from_query(user_query)


def _tool_input_for_request(user_query: str, tool_name: str, analysis: dict | None) -> str:
    if tool_name in _TOOLS_KEEP_FULL_QUERY:
        return user_query
    if tool_name in _TOOLS_GLOBAL_SCOPE:
        return ""
    if not (user_query or "").strip():
        return ""
    if analysis is not None:
        return _monitoring_tool_input_from_analysis(user_query, analysis)
    return extract_service_name_from_query(user_query)


# ✅ Flask App
flask_app = Flask(__name__, template_folder='templates')
flask_app.secret_key = (
    os.getenv("FLASK_SECRET_KEY")
    or os.getenv("ADMIN_TOKEN")
    or os.urandom(32)
)
CORS(flask_app, supports_credentials=True)


# Unified message: Docker image omits .env (.dockerignore); inject vars or mount .env on the host.
_SLACK_WEBHOOK_MISSING_MSG = (
    "SLACK_WEBHOOK_URL is not set. "
    "The Docker image does not include .env: set SLACK_WEBHOOK_URL on the host and run "
    "docker run ... --env-file .env, or use docker-compose (env_file: .env). "
    "See DOCKER_DEPLOYMENT.md for required variables."
)

# Short-lived PNG cache for Slack (webhook has no multipart; Slack fetches image_url)
_SLACK_SCREENSHOT_CACHE = {}
_SLACK_SCREENSHOT_LOCK = threading.Lock()
_SLACK_SCREENSHOT_TTL_SEC = 300


def _slack_screenshot_prune_locked():
    now = time.time()
    dead = [k for k, (_, exp) in _SLACK_SCREENSHOT_CACHE.items() if exp < now]
    for k in dead:
        del _SLACK_SCREENSHOT_CACHE[k]


def _slack_screenshot_store_png(data: bytes) -> str:
    sid = uuid.uuid4().hex
    exp = time.time() + _SLACK_SCREENSHOT_TTL_SEC
    with _SLACK_SCREENSHOT_LOCK:
        _slack_screenshot_prune_locked()
        _SLACK_SCREENSHOT_CACHE[sid] = (data, exp)
    return sid


def _slack_screenshot_get_png(sid: str):
    with _SLACK_SCREENSHOT_LOCK:
        _slack_screenshot_prune_locked()
        t = _SLACK_SCREENSHOT_CACHE.get(sid)
        if not t:
            return None
        data, exp = t
        if time.time() > exp:
            del _SLACK_SCREENSHOT_CACHE[sid]
            return None
        return data


def _slack_screenshot_public_base_url():
    """HTTPS base URL where Slack can download the PNG (this app reachable from the internet)."""
    base = (
        os.getenv("PUBLIC_BASE_URL")
        or os.getenv("SLACK_IMAGE_PUBLIC_BASE_URL")
        or ""
    ).strip().rstrip("/")
    if base:
        return base
    return request.url_root.rstrip("/")


alerts_db = []  # {id, text, priority, ack, cause}

def classify_alert(text):
    text_lower = text.lower()
    if any(k in text_lower for k in ['Sev1', 'Sev0']):
        return 'High'
    if any(k in text_lower for k in ['Sev3', 'Sev2']):
        return 'Medium'
    return 'Low'


def identify_cause(text):
    try:
        # Run real tools with the alert text
        # wiki_info = TOOLS["Wiki"]["function"](text)
        confluence_info = TOOLS["Wiki"]["function"](text)
        suggestion = TOOLS["Suggestions"]["function"](text)

        return f""" 
        <div>
            <h4>Suggested root cause:</h4>{suggestion}
            <br><h4>Confluence:</h4>{confluence_info}
        </div>
        """
    except Exception as e:
        return f"<pre>Error identifying root cause: {e}</pre>"



@flask_app.route('/')
def index():
    return render_template('index.html')

@flask_app.route('/api/history')
def api_history():
    return jsonify(get_history())


# Status Monitor Dashboard Routes
@flask_app.route('/statusmonitor')
def statusmonitor_page():
    """Hub: overview of each environment; click through to /statusmonitor/<env>."""
    return render_template('statusmonitor.html', environment=None)


@flask_app.route('/statusmonitor/production')
def statusmonitor_production_page():
    """Serve the status monitor dashboard page for production only"""
    return render_template('statusmonitor.html', environment='production')


@flask_app.route('/statusmonitor/goldendev')
def statusmonitor_goldendev_page():
    """Serve the status monitor dashboard page for goldendev only"""
    return render_template('statusmonitor.html', environment='goldendev')


@flask_app.route('/statusmonitor/goldenqa')
def statusmonitor_goldenqa_page():
    """Serve the status monitor dashboard page for goldenqa only"""
    return render_template('statusmonitor.html', environment='goldenqa')


@flask_app.route('/statusmonitor/qa')
def statusmonitor_qa_page():
    """Serve the status monitor dashboard for env:qa (cluster/platform list)"""
    return render_template('statusmonitor.html', environment='qa')


@flask_app.route('/statusmonitor/samsung')
def statusmonitor_samsung_page():
    """Serve the status monitor dashboard page for Samsung network services only"""
    return render_template('statusmonitor.html', environment='samsung')


@flask_app.route('/statusmonitor/adt')
def statusmonitor_adt_page():
    """Serve the status monitor dashboard page for ADT partner services only"""
    return render_template('statusmonitor.html', environment='adt')


@flask_app.route('/statusmonitor/cat')
def statusmonitor_cat_page():
    """Serve the status monitor dashboard page for CAT partner services only"""
    return render_template('statusmonitor.html', environment='cat')


@flask_app.route('/statusmonitor/comcast')
def statusmonitor_comcast_page():
    """Serve the status monitor dashboard page for Comcast partner services only"""
    return render_template('statusmonitor.html', environment='comcast')


@flask_app.route('/statusmonitor/redmetrics-us')
def statusmonitor_redmetrics_us_page():
    """Serve the status monitor dashboard page for RED Metrics US services"""
    return render_template('statusmonitor.html', environment='redmetrics-us')


@flask_app.route("/embed/splunk-p0-adt")
def embed_splunk_p0_adt():
    """P0 ADT Streaming charts (Splunk REST) for ADT Status Wall in-page panel."""
    from tools.splunk_tool import read_splunk_p0_adt_dashboard

    tr = request.args.get("timerange", type=int)
    service = (request.args.get("service") or request.args.get("host") or "").strip()
    body = read_splunk_p0_adt_dashboard(service, timerange=tr)
    return (
        "<!DOCTYPE html><html lang=\"en\"><head><meta charset=\"utf-8\">"
        "<meta name=\"viewport\" content=\"width=device-width,initial-scale=1\">"
        "<title>Splunk P0 ADT</title>"
        "<style>body{margin:0;padding:12px 14px 24px;background:#0b0c12;color:#e2e8f0;"
        "font-family:system-ui,sans-serif;} a{color:#7dd3fc;}</style></head><body>"
        f"{body}</body></html>"
    )


@flask_app.route("/embed/splunk-samsung-latencies")
def embed_splunk_samsung_latencies():
    """
    Samsung alarm-latency charts from Splunk REST (port 8089) + token — same stack as
    the Samsung Dashboard project (samsung_splunk_api_latencies). No Splunk web login in browser.
    """
    from samsung_splunk_api_latencies import build_embed_for_flask

    se = (os.environ.get("SPLUNK_EMBED_EARLIEST") or "-30d@d").strip()
    sl = (os.environ.get("SPLUNK_EMBED_LATEST") or "now").strip()
    try:
        leg = int((os.environ.get("SPLUNK_EMBED_LEGACY_HOURS") or "720").strip() or "720")
    except ValueError:
        leg = 720
    return build_embed_for_flask(leg, studio_earliest=se, studio_latest=sl)


@flask_app.route("/aws-change-tracker")
def aws_change_tracker_page():
    """AWS operations UI: CloudTrail change lookup + Amazon Connect monitoring."""
    from tools.aws_cloudtrail_tracker import CLOUDTRAIL_RESOURCE_TYPE_OPTIONS
    from tools.aws_connect_monitor import connect_monitor_config

    return render_template(
        "aws_change_tracker.html",
        resource_type_options=CLOUDTRAIL_RESOURCE_TYPE_OPTIONS,
        connect_config=connect_monitor_config(),
    )


@flask_app.route("/api/aws/cloudtrail/search", methods=["POST"])
def api_aws_cloudtrail_search():
    """JSON: CloudTrail LookupEvents filtered by resource name + account + optional resource type."""
    try:
        from tools.aws_cloudtrail_tracker import cloudtrail_search

        data = request.get_json() or {}
        raw_max = data.get("max_events")
        if raw_max is None or str(raw_max).strip() == "":
            max_ev = 50
        else:
            try:
                max_ev = int(raw_max)
            except (TypeError, ValueError):
                max_ev = 50
        try:
            lb = int(data.get("lookback_days") or 7)
        except (TypeError, ValueError):
            lb = 7
        out = cloudtrail_search(
            resource_name=str(data.get("resource_name") or ""),
            resource_type=str(data.get("resource_type") or ""),
            region=str(data.get("region") or ""),
            account_id=str(data.get("account_id") or ""),
            lookback_days=lb,
            max_events=max_ev,
        )
        if not out.get("success"):
            return jsonify(out), 400
        return jsonify(out)
    except Exception as e:
        logging.exception("CloudTrail search API")
        return jsonify({"success": False, "error": str(e)}), 500


@flask_app.route("/api/aws/cloudtrail/analyze-upload", methods=["POST"])
def api_aws_cloudtrail_analyze_upload():
    """Parse a console-exported CSV (UTF-8); no AWS calls."""
    try:
        from tools.aws_cloudtrail_tracker import parse_console_csv_or_excel

        f = request.files.get("file")
        if not f or not f.filename:
            return jsonify({"success": False, "error": "Missing form field 'file'."}), 400
        out = parse_console_csv_or_excel(f)
        if not out.get("success"):
            return jsonify(out), 400
        return jsonify(out)
    except Exception as e:
        logging.exception("CloudTrail CSV analyze")
        return jsonify({"success": False, "error": str(e)}), 500


@flask_app.route("/api/aws/connect/monitor", methods=["GET", "POST"])
def api_aws_connect_monitor():
    """Amazon Connect operational snapshot: critical alerts, proactive checks, dashboard."""
    try:
        from tools.aws_connect_monitor import connect_monitor_snapshot

        data = request.get_json(silent=True) or {}
        if request.method == "GET":
            data = request.args
        force = str(data.get("force_refresh") or "").lower() in ("1", "true", "yes")
        out = connect_monitor_snapshot(
            instance_id=str(data.get("instance_id") or "").strip() or None,
            region=str(data.get("region") or "").strip() or None,
            force_refresh=force,
        )
        if not out.get("success"):
            return jsonify(out), 400
        return jsonify(out)
    except Exception as e:
        logging.exception("Connect monitor API")
        return jsonify({"success": False, "error": str(e)}), 500


@flask_app.route('/apm-services')
def apm_services_page():
    """
    APM Status Wall: default `all` = every APM `env` as its own block; or one
    (production, goldendev, goldenqa, adt_prod, samsung_prod, cat_prod, comcast_prod) via `?dd_env=…`.
    The `qa` env is not part of the Main tab aggregate or dropdown; `dd_env=qa` still works for a focused API/query.
    See SOFTWARE_CATALOG_* and lists/*_apm_services.txt.
    """
    from tools.status_monitor import (
        SOFTWARE_CATALOG_WALL_APM_ENVS,
        SOFTWARE_CATALOG_WALL_GOLDEN_ENVS,
        _apm_wall_group_label,
        normalize_software_catalog_wall_dd_env,
    )
    import re

    tab = (request.args.get("tab") or "").strip().lower()
    q_raw = (request.args.get("dd_env") or os.environ.get("APM_STATUS_WALL_DD_ENV") or "production").strip()
    if tab == "golden":
        wall_dd_env = normalize_software_catalog_wall_dd_env("golden")
    else:
        wall_dd_env = normalize_software_catalog_wall_dd_env(q_raw)
    # Datadog Software list uses one `env` tag; when showing all, link to a neutral default.
    if wall_dd_env == "all":
        _dd_sw = "production"
    elif wall_dd_env == "golden":
        _dd_sw = "goldendev"
    else:
        _dd_sw = wall_dd_env
    dd_base = (os.environ.get("DATADOG_APM_SOFTWARE_BASE") or "").strip()
    if not dd_base:
        datadog_software_href = (
            f"https://arlo.datadoghq.com/software?env={_dd_sw}&fromUser=true"
        )
    else:
        if re.search(r"[?&]env=", dd_base):
            datadog_software_href = re.sub(
                r"env=[^&]*", f"env={_dd_sw}", dd_base, count=1
            )
        else:
            sep = "&" if "?" in dd_base else "?"
            datadog_software_href = f"{dd_base}{sep}env={_dd_sw}"
            if "fromUser" not in datadog_software_href:
                qm = "?" if "?" not in datadog_software_href else "&"
                datadog_software_href = f"{datadog_software_href}{qm}fromUser=true"
    if wall_dd_env == "golden":
        _slack = "Status Wall — Golden (goldendev + goldenqa)"
    elif wall_dd_env == "all":
        _slack = "Status Wall — all envs"
    elif wall_dd_env != "production":
        _slack = f"Status Wall — {wall_dd_env}"
    else:
        _slack = "Status Wall — production"
    wall_apm_tab = "golden" if wall_dd_env == "golden" else "main"
    wall_apm_env_labels = {
        e: _apm_wall_group_label(e)
        for e in tuple(SOFTWARE_CATALOG_WALL_APM_ENVS)
        + tuple(SOFTWARE_CATALOG_WALL_GOLDEN_ENVS)
    }
    _wall_title = "Status Wall"
    resp = make_response(
        render_template(
        "statuswall.html",
        wall_title=_wall_title,
        wall_api="/api/statusmonitor/software-catalog-wall",
        wall_nav="apm_wall",
        wall_slack_title=_slack,
        wall_show_apm_env=True,
        wall_adt_splunk_embed=(wall_dd_env == "adt_prod"),
        wall_dd_env=wall_dd_env,
        wall_apm_tab=wall_apm_tab,
        datadog_software_href=datadog_software_href,
        wall_incremental_apm=True,
        wall_apm_parallel_main_envs=list(SOFTWARE_CATALOG_WALL_APM_ENVS),
        wall_apm_parallel_golden_envs=list(SOFTWARE_CATALOG_WALL_GOLDEN_ENVS),
        wall_apm_env_labels=wall_apm_env_labels,
        )
    )
    resp.headers["Cache-Control"] = "no-store, max-age=0"
    return resp


@flask_app.route('/api/statusmonitor', methods=['POST'])
def api_statusmonitor():
    """API endpoint for status monitor dashboard data"""
    try:
        from tools.status_monitor import status_monitor_dashboard
        
        data = request.get_json() or {}
        timerange = data.get('timerange', 1)
        environment = data.get('environment', None)  # Optional: specific environment
        force_refresh = bool(data.get('force_refresh') or data.get('forceRefresh'))
        
        html_content = status_monitor_dashboard(
            timerange=timerange, environment=environment, force_refresh=force_refresh
        )
        
        return jsonify({
            'success': True,
            'html': html_content
        })
    except Exception as e:
        logging.error(f"Error in status monitor: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@flask_app.route('/api/statusmonitor/partial', methods=['POST'])
def api_statusmonitor_partial():
    """Incremental fragments for /statusmonitor/<env> (bootstrap, sidebar, per-env APM, finalize)."""
    try:
        from tools.status_monitor import status_monitor_partial

        data = request.get_json() or {}
        timerange = int(data.get('timerange', 1))
        environment = data.get('environment')
        force_refresh = bool(data.get('force_refresh') or data.get('forceRefresh'))
        part = data.get('part') or 'meta'
        session_id = data.get('session_id') or data.get('sessionId')
        dd_env = data.get('dd_env') or data.get('ddEnv')
        payload = status_monitor_partial(
            part,
            timerange=timerange,
            environment=environment,
            force_refresh=force_refresh,
            session_id=session_id,
            dd_env=dd_env,
        )
        if not payload.get('success'):
            return jsonify(payload), 400
        return jsonify(payload)
    except Exception as e:
        logging.error(f"Error in status monitor partial ({part}): {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e), 'part': (request.get_json(silent=True) or {}).get('part')}), 500


@flask_app.route('/api/statusmonitor/hub-summary', methods=['POST'])
def api_statusmonitor_hub_summary():
    """JSON summary for /statusmonitor hub (one row per environment, Datadog health)."""
    try:
        from tools.status_monitor import status_monitor_hub_summary

        data = request.get_json() or {}
        timerange = int(data.get('timerange', 1))
        force_refresh = bool(data.get('force_refresh') or data.get('forceRefresh'))
        return jsonify(status_monitor_hub_summary(timerange=timerange, force_refresh=force_refresh))
    except Exception as e:
        logging.error(f"Error in status monitor hub summary: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@flask_app.route('/api/statusmonitor/wall', methods=['POST'])
def api_statusmonitor_wall():
    """JSON payload for /statuswall (grouped service tiles, same health logic as hub)."""
    try:
        from tools.status_monitor import status_monitor_wall_data

        data = request.get_json() or {}
        timerange = int(data.get('timerange', 1))
        force_refresh = bool(data.get('force_refresh') or data.get('forceRefresh'))
        return jsonify(status_monitor_wall_data(timerange=timerange, force_refresh=force_refresh))
    except Exception as e:
        logging.error(f"Error in status monitor wall: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@flask_app.route("/api/statusmonitor/software-catalog-wall", methods=["POST"])
def api_statusmonitor_software_catalog_wall():
    """
    JSON for /apm-services: APM Status Wall. Body: dd_env
    = all (default) for main envs, golden for Golden tab (goldendev + goldenqa),
    or one of production|goldendev|goldenqa|adt_prod|samsung_prod|cat_prod|comcast_prod|qa (qa not in Main “all” list).
    """
    try:
        from tools.status_monitor import (
            normalize_software_catalog_wall_dd_env,
            status_monitor_software_catalog_wall_data,
        )

        data = request.get_json() or {}
        timerange = int(data.get("timerange", 24))
        force_refresh = bool(
            data.get("force_refresh") or data.get("forceRefresh")
        )
        raw_dd = data.get("dd_env") or data.get("ddEnv")
        dd_e = normalize_software_catalog_wall_dd_env(
            (raw_dd if raw_dd is not None and str(raw_dd).strip() else None)
        )
        return jsonify(
            status_monitor_software_catalog_wall_data(
                timerange=timerange, force_refresh=force_refresh, dd_env=dd_e
            )
        )
    except Exception as e:
        logging.error(f"Error in software catalog wall: {e}")
        import traceback

        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


# ========================================
# REST API Endpoints for Metrics & History
# ========================================

@flask_app.route('/api/status/current', methods=['GET'])
def api_status_current():
    """
    Get current status for all services (JSON format)
    Query params:
        - environment: Filter by environment (optional)
    """
    try:
        from tools.metrics_persistence import get_all_services_current_status
        
        environment = request.args.get('environment')
        services = get_all_services_current_status()
        
        # Filter by environment if specified
        if environment:
            services = [s for s in services if s.get('environment') == environment]
        
        return jsonify({
            'success': True,
            'timestamp': datetime.utcnow().isoformat(),
            'total_services': len(services),
            'services': services
        })
    except Exception as e:
        logging.error(f"Error fetching current status: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@flask_app.route('/api/status/<environment>', methods=['GET'])
def api_status_by_environment(environment):
    """
    Get current status for specific environment (JSON format)
    Path params:
        - environment: production, goldendev, or goldenqa
    """
    try:
        from tools.metrics_persistence import get_all_services_current_status
        
        services = get_all_services_current_status()
        env_services = [s for s in services if s.get('environment') == environment]
        
        # Calculate summary
        total = len(env_services)
        healthy = sum(1 for s in env_services if s['status'] == 'healthy')
        warning = sum(1 for s in env_services if s['status'] == 'warning')
        critical = sum(1 for s in env_services if s['status'] == 'critical')
        
        return jsonify({
            'success': True,
            'timestamp': datetime.utcnow().isoformat(),
            'environment': environment,
            'summary': {
                'total_services': total,
                'healthy': healthy,
                'warning': warning,
                'critical': critical
            },
            'services': env_services
        })
    except Exception as e:
        logging.error(f"Error fetching status for {environment}: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@flask_app.route('/api/history/service/<service_name>', methods=['GET'])
def api_service_history(service_name):
    """
    Get historical metrics for a specific service
    Path params:
        - service_name: Name of the service
    Query params:
        - environment: Environment (required)
        - hours: Hours to look back (default: 24)
    """
    try:
        from tools.metrics_persistence import get_service_history
        
        environment = request.args.get('environment')
        if not environment:
            return jsonify({
                'success': False,
                'error': 'environment parameter is required'
            }), 400
        
        hours = int(request.args.get('hours', 24))
        history = get_service_history(service_name, environment, hours)
        
        return jsonify({
            'success': True,
            'service': service_name,
            'environment': environment,
            'hours': hours,
            'data_points': len(history),
            'history': history
        })
    except Exception as e:
        logging.error(f"Error fetching service history: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@flask_app.route('/api/history/dashboard', methods=['GET'])
def api_dashboard_history():
    """
    Get historical dashboard snapshots
    Query params:
        - environment: Filter by environment (optional)
        - hours: Hours to look back (default: 24)
    """
    try:
        from tools.metrics_persistence import get_dashboard_history
        
        environment = request.args.get('environment')
        hours = int(request.args.get('hours', 24))
        
        history = get_dashboard_history(environment, hours)
        
        return jsonify({
            'success': True,
            'environment': environment or 'all',
            'hours': hours,
            'data_points': len(history),
            'history': history
        })
    except Exception as e:
        logging.error(f"Error fetching dashboard history: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@flask_app.route('/api/trends/service/<service_name>', methods=['GET'])
def api_service_trends(service_name):
    """
    Get trend analysis for a specific service
    Path params:
        - service_name: Name of the service
    Query params:
        - environment: Environment (required)
        - hours: Hours to analyze (default: 24)
    """
    try:
        from tools.metrics_persistence import get_service_trends
        
        environment = request.args.get('environment')
        if not environment:
            return jsonify({
                'success': False,
                'error': 'environment parameter is required'
            }), 400
        
        hours = int(request.args.get('hours', 24))
        trends = get_service_trends(service_name, environment, hours)
        
        return jsonify({
            'success': True,
            'trends': trends
        })
    except Exception as e:
        logging.error(f"Error calculating trends: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@flask_app.route('/api/critical/history', methods=['GET'])
def api_critical_history():
    """
    Get history of critical service incidents
    Query params:
        - hours: Hours to look back (default: 24)
    """
    try:
        from tools.metrics_persistence import get_critical_services_history
        
        hours = int(request.args.get('hours', 24))
        critical_history = get_critical_services_history(hours)
        
        return jsonify({
            'success': True,
            'hours': hours,
            'total_incidents': len(critical_history),
            'incidents': critical_history
        })
    except Exception as e:
        logging.error(f"Error fetching critical history: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@flask_app.route('/api/health', methods=['GET'])
def api_health():
    """
    Liveness-style check for load balancers and Docker HEALTHCHECK.

    Always returns HTTP 200 if this worker can handle requests. SQLite metrics
    are optional: failures are reported as status 'degraded' so ALB/Ingress
    does not drop the target (which would surface as 503 for every route).
    """
    try:
        from tools.metrics_persistence import get_database_stats

        db_stats = get_database_stats()
        return jsonify({
            'status': 'healthy',
            'timestamp': datetime.utcnow().isoformat(),
            'database': db_stats,
            'version': '3.0.2'
        })
    except Exception as e:
        logging.warning("Health check: metrics DB unavailable: %s", e)
        return jsonify({
            'status': 'degraded',
            'timestamp': datetime.utcnow().isoformat(),
            'database': {'error': str(e)},
            'version': '3.0.2'
        })


@flask_app.route('/testconnections')
def testconnections_page():
    """Página para comprobar integraciones (lee .env del proceso; no expone secretos)."""
    return render_template('testconnections.html')


@flask_app.route('/secrets')
def secrets_page():
    """Formulario para editar tokens del .env (valores enmascarados)."""
    from tools.dev_admin import DEFAULT_SECRETS_PIN

    return render_template('secrets.html', secrets_pin=DEFAULT_SECRETS_PIN)


@flask_app.route('/api/secrets', methods=['GET'])
def api_secrets_get():
    try:
        from tools.env_secrets import list_secrets_state

        return jsonify(list_secrets_state())
    except Exception as e:
        logging.exception('api_secrets_get failed')
        return jsonify({'error': str(e)}), 500


@flask_app.route('/api/secrets', methods=['POST'])
def api_secrets_save():
    from tools import dev_admin as dev_admin_mod
    from tools.env_secrets import update_secrets_in_dotenv

    ok, msg = dev_admin_mod.verify_secrets_save_pin(_admin_token_from_request())
    if not ok:
        return jsonify({'success': False, 'error': msg}), 401
    try:
        body = request.get_json(silent=True) or {}
        updates = body.get('updates') if isinstance(body.get('updates'), dict) else body
        if not isinstance(updates, dict):
            return jsonify({'success': False, 'error': 'JSON inválido: usa {"updates": {"KEY": "valor"}}.'}), 400
        out = update_secrets_in_dotenv(updates)
        status = 200 if out.get('success') else 400
        return jsonify(out), status
    except Exception as e:
        logging.exception('api_secrets_save failed')
        return jsonify({'success': False, 'error': str(e)}), 500


@flask_app.route('/api/testconnections', methods=['GET', 'POST'])
def api_testconnections():
    try:
        from tools.test_connections import run_connection_checks

        return jsonify(run_connection_checks())
    except Exception as e:
        logging.exception('api_testconnections failed')
        return jsonify({'error': str(e), 'items': [], 'all_ok': False}), 500


def _admin_token_from_request():
    t = (request.headers.get('X-Admin-Token') or request.form.get('token') or '').strip()
    if t:
        return t
    if request.is_json:
        return str((request.get_json(silent=True) or {}).get('token') or '').strip()
    return ''


@flask_app.route('/dev-admin')
def dev_admin_page():
    """Git pull + reemplazo de .env (APIs protegidas con ADMIN_TOKEN)."""
    from tools import dev_admin as dev_admin_mod

    return render_template('dev_admin.html', admin_enabled=dev_admin_mod.admin_token_configured())


@flask_app.route('/api/admin/git-update', methods=['POST'])
def api_admin_git_update():
    from tools import dev_admin as dev_admin_mod

    ok, msg = dev_admin_mod.verify_admin_request_token(_admin_token_from_request())
    if not ok:
        code = 403 if 'deshabilitad' in msg.lower() or 'no está definido' in msg.lower() else 401
        return jsonify({'success': False, 'error': msg}), code
    try:
        out = dev_admin_mod.git_update_status_and_pull()
        status = 200 if out.get('success') else 400
        return jsonify(out), status
    except Exception as e:
        logging.exception('api_admin_git_update')
        return jsonify({'success': False, 'error': str(e)}), 500


@flask_app.route('/api/admin/upload-env', methods=['POST'])
def api_admin_upload_env():
    from tools import dev_admin as dev_admin_mod

    ok, msg = dev_admin_mod.verify_admin_request_token(_admin_token_from_request())
    if not ok:
        code = 403 if 'deshabilitad' in msg.lower() or 'no está definido' in msg.lower() else 401
        return jsonify({'success': False, 'error': msg}), code
    try:
        f = request.files.get('file')
        if f and f.filename:
            out = dev_admin_mod.save_uploaded_dotenv(f)
        elif request.is_json:
            body = request.get_json(silent=True) or {}
            out = dev_admin_mod.save_dotenv_text_content(body.get('content'))
        else:
            out = {'success': False, 'error': 'Envía un archivo (campo file) o JSON {"content": "..."}.'}
        status = 200 if out.get('success') else 400
        return jsonify(out), status
    except Exception as e:
        logging.exception('api_admin_upload_env')
        return jsonify({'success': False, 'error': str(e)}), 500


@flask_app.route('/api/cache/clear', methods=['POST'])
def api_clear_cache():
    """Clear the status monitor cache (force fresh data on next load)"""
    try:
        from tools.status_monitor import clear_status_cache
        
        clear_status_cache()
        
        return jsonify({
            'success': True,
            'message': 'Cache cleared successfully',
            'timestamp': datetime.utcnow().isoformat()
        })
    except Exception as e:
        logging.error(f"Error clearing cache: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@flask_app.route('/api/tools')
def api_tools():
    return jsonify([{'name': name, 'desc': desc} for name, desc in registered_tools])


@flask_app.route('/api/mcp/tools')
def api_mcp_tools():
    """Unified categorized tools for checkbox UI (MCP local + synthesis-only legacy)."""
    from mcp_server import TOOL_REGISTRY
    from tools.mcp_tool_catalog import UI_SYNTHESIS_TOOL_NAMES, build_ui_tool_catalog

    synthesis = {
        name: TOOLS[name]["description"]
        for name in UI_SYNTHESIS_TOOL_NAMES
        if name in TOOLS
    }
    return jsonify(build_ui_tool_catalog(TOOL_REGISTRY, synthesis_tools=synthesis))


@flask_app.route('/api/suggest-tools', methods=['POST'])
def suggest_tools():
    """Use AI to suggest which tools to use based on the user's query"""
    data = request.get_json()
    user_query = data.get('query', '').strip()
    
    if not user_query:
        return jsonify({'error': 'No query provided'}), 400
    
    try:
        logging.info(f"🤖 AI Auto-Select: Analyzing query: {user_query[:100]}")
        
        from mcp_server import TOOL_REGISTRY
        from tools.mcp_tool_suggest import (
            augment_suggested_tools_for_query,
            build_suggest_tools_catalog_text,
            is_service_health_question,
            normalize_and_validate_suggested_tools,
            service_health_mcp_checkboxes,
        )
        from tools.service_query import extract_service_name_from_query
        from tools.shm_tools import is_shm_metrics_question

        available_tools = build_suggest_tools_catalog_text(TOOL_REGISTRY)
        service_hint = ""
        if is_service_health_question(user_query):
            svc = extract_service_name_from_query(user_query)
            service_hint = (
                f'\nDetected service-specific health query for "{svc}" — MUST include '
                "MCP:datadog_services, MCP:datadog_search, MCP:datadog_errors, MCP:datadog_red_metrics, "
                "MCP:service_owners, and Bedrock_Report.\n"
            )
        elif is_shm_metrics_question(user_query):
            service_hint = (
                "\nDetected SHM / customer satisfaction query — MUST include MCP:shm_metrics "
                "(iOS/Android app ratings, CSAT pillar, Chart.js graphs from shmview.arlocloud.com). "
                "Do NOT use MintMCP Amplitude tools for this. Include Bedrock_Report.\n"
            )

        analysis_prompt = f"""Analyze this question and select the appropriate tools.

QUESTION: "{user_query}"
{service_hint}
AVAILABLE TOOLS (use exact checkbox values):
{available_tools}

RULES:
1. Return checkbox values exactly as listed (e.g. MCP:datadog_services, Bedrock_Report).
2. For a SPECIFIC SERVICE (errors, status, what is wrong, inappayments, backend-hmsfoo): include MCP:datadog_services, MCP:datadog_search, MCP:datadog_errors, MCP:datadog_red_metrics, MCP:service_owners, and Bedrock_Report.
3. For SHM / Service Health Management, customer satisfaction, pillar scores, iOS/Android app metrics: MCP:shm_metrics and/or MCP:shm_daily + Bedrock_Report.
4. For incidents/alerts: add MCP:pagerduty_incidents.
5. For Confluence/docs only: MCP:wiki_search + Bedrock_Report.
6. Bedrock_Report synthesizes all data — include it for any data lookup (not pure definitions).
7. Select ALL relevant tools; err on the side of more Datadog tools for service questions.

Return ONLY a JSON array: ["MCP:datadog_services", "MCP:datadog_errors", "Bedrock_Report"]
NO markdown, NO explanation."""

        # Call Bedrock to get tool suggestions
        logging.info("🤖 Calling Bedrock for tool recommendations...")
        suggested_tools_response = ask_bedrock(analysis_prompt, selected_tools=None)
        logging.info(f"🤖 Bedrock response: {suggested_tools_response[:200]}")
        
        # Parse the JSON response
        # Try to extract JSON array from response
        json_match = re.search(r'\[.*?\]', suggested_tools_response, re.DOTALL)
        if json_match:
            suggested_tools = json.loads(json_match.group(0))
        else:
            # Fallback: if no JSON found, return error
            logging.error(f"❌ Could not parse Bedrock response: {suggested_tools_response}")
            return jsonify({'error': 'Could not parse AI response', 'raw_response': suggested_tools_response}), 500
        
        # Validate and map to UI checkbox values (MCP + synthesis)
        valid_tools = normalize_and_validate_suggested_tools(suggested_tools)
        valid_tools = augment_suggested_tools_for_query(user_query, valid_tools)

        if is_service_health_question(user_query) and not any(
            t.startswith("MCP:datadog_") for t in valid_tools
        ):
            for cb in service_health_mcp_checkboxes():
                if cb not in valid_tools:
                    valid_tools.append(cb)
            logging.info("➕ Heuristic: added Datadog MCP tools for service health query")
        
        # 🔥 ALWAYS ADD Bedrock_Report for comprehensive context (unless it's pure explanation)
        # Bedrock_Report executes LAST but displays FIRST (for better UX)
        user_query_lower = user_query.lower()
        
        # Check if this is a pure explanation query (no data lookup needed)
        pure_explanation_keywords = ['what is', 'qué es', 'que es', 'explain', 'explica', 'define']
        is_pure_explanation = (
            any(keyword in user_query_lower for keyword in pure_explanation_keywords) and
            len(valid_tools) == 1 and 
            'Ask_Bedrock' in valid_tools
        )
        
        if not is_pure_explanation and 'Bedrock_Report' not in valid_tools:
            valid_tools.append('Bedrock_Report')
            logging.info(f"➕ Auto-adding Bedrock_Report for comprehensive context synthesis")
        
        # Reorder: Put Bedrock_Report FIRST for display (it will still execute last due to phase logic)
        if 'Bedrock_Report' in valid_tools:
            valid_tools.remove('Bedrock_Report')
            valid_tools.insert(0, 'Bedrock_Report')  # Insert at beginning for UI
            logging.info(f"🔄 Moved Bedrock_Report to FIRST position for UI display")
        
        logging.info(f"✅ Final tool selection: {len(valid_tools)} tool(s): {valid_tools}")
        return jsonify({'suggested_tools': valid_tools})
        
    except Exception as e:
        logging.error(f"❌ Error in suggest-tools: {e}")
        return jsonify({'error': str(e)}), 500

@flask_app.route('/api/run', methods=['POST'])
def api_run():
    data = request.json
    input_text = data.get('input', '')
    selected_tools = list(data.get('tools', []) or ['Suggestions'])
    timerange = data.get('timerange', 4)  # Default to 4 hours
    pagerduty_filters = {
        "shift": str(data.get("pagerduty_shift") or "").strip().lower(),
        "team_only": bool(data.get("pagerduty_team_only")),
        "missing_root_cause": bool(data.get("pagerduty_missing_rca")),
    }
    start = time.time()

    from tools.mcp_tool_suggest import augment_suggested_tools_for_query, is_service_health_question
    from tools.shm_tools import is_shm_daily_question, is_shm_metrics_question

    synthesis_tools_set = frozenset({'Bedrock_Report', 'Ask_Bedrock'})
    needs_auto_data_tools = (
        is_service_health_question(input_text)
        or is_shm_metrics_question(input_text)
        or is_shm_daily_question(input_text)
    )
    if input_text.strip() and needs_auto_data_tools:
        data_only = [t for t in selected_tools if t not in synthesis_tools_set]
        if not data_only:
            augmented = augment_suggested_tools_for_query(input_text, selected_tools)
            if augmented != selected_tools:
                logging.info(
                    "➕ Auto-injecting MCP data tools (only synthesis was selected): %s",
                    [t for t in augmented if t not in selected_tools],
                )
                selected_tools = augmented
    
    # Execute tools in parallel using threading
    import concurrent.futures
    from threading import Lock
    
    results_dict = {}
    results_lock = Lock()
    
    def analyze_query_with_bedrock(user_query: str) -> dict:
        """
        Use Bedrock to intelligently analyze user query and extract intent
        Returns: {
            'is_general_query': bool,  # True if asking for all services
            'service_name': str,        # Empty if general, or specific service name
            'confidence': str           # 'high', 'medium', 'low'
        }
        """
        try:
            bedrock_runtime = _get_bedrock_runtime_client()
            if not bedrock_runtime:
                logging.warning("⚠️ BEDROCK_API_KEY not available for query analysis")
                return {
                    "is_general_query": False,
                    "service_name": extract_service_name_from_query(user_query),
                    "confidence": "low",
                }
            
            analysis_prompt = f"""Analyze this user query and extract the intent for monitoring tool execution.

User Query: "{user_query}"

Determine:
1. Is this a GENERAL query asking for ALL services/dashboards? (e.g., "all services", "all zones", "general status", "show everything", "red metrics for all regions")
2. Or is it asking for a SPECIFIC service? (e.g., "hmsguard status", "backend-hmsalexaapi metrics", "device-location errors")

If SPECIFIC, extract the exact service name (e.g., "hmsguard", "backend-hmsalexaapi", "device-location").

Respond ONLY with valid JSON (no markdown, no explanations):
{{
    "is_general_query": true/false,
    "service_name": "extracted-service-name or empty string",
    "confidence": "high/medium/low"
}}

Examples:
- "show me red metrics for all zones" → {{"is_general_query": true, "service_name": "", "confidence": "high"}}
- "what's happening with hmsguard?" → {{"is_general_query": false, "service_name": "hmsguard", "confidence": "high"}}
- "give me results for all regions" → {{"is_general_query": true, "service_name": "", "confidence": "high"}}
- "backend-hmsalexaapi errors" → {{"is_general_query": false, "service_name": "backend-hmsalexaapi", "confidence": "high"}}"""

            request_body = {
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 500,
                "temperature": 0.1,
                "messages": [{"role": "user", "content": analysis_prompt}]
            }
            
            response = bedrock_runtime.invoke_model(
                modelId="us.anthropic.claude-sonnet-4-20250514-v1:0",
                body=json.dumps(request_body)
            )
            
            response_body = json.loads(response['body'].read())
            bedrock_response = response_body.get('content', [{}])[0].get('text', '{}')
            
            # Parse JSON response
            # Remove markdown code blocks if present
            bedrock_response = bedrock_response.strip()
            if bedrock_response.startswith('```'):
                bedrock_response = bedrock_response.split('```')[1]
                if bedrock_response.startswith('json'):
                    bedrock_response = bedrock_response[4:]
                bedrock_response = bedrock_response.strip()
            
            analysis = json.loads(bedrock_response)
            
            logging.info(f"🤖 Bedrock Query Analysis:")
            logging.info(f"   - User Query: '{user_query}'")
            logging.info(f"   - Is General: {analysis.get('is_general_query', False)}")
            logging.info(f"   - Service Name: '{analysis.get('service_name', '')}'")
            logging.info(f"   - Confidence: {analysis.get('confidence', 'unknown')}")
            
            return analysis
            
        except Exception as e:
            logging.error(f"❌ Error in Bedrock query analysis: {e}")
            return {
                "is_general_query": False,
                "service_name": extract_service_name_from_query(user_query),
                "confidence": "low",
            }
    
    def execute_tool(idx, tool_name, context_from_other_tools=None, query_analysis=None):
        """Execute a single tool and store result."""
        from tools.mcp_tool_catalog import build_mcp_tool_arguments, parse_mcp_checkbox_value
        from tools.mcp_tool_dispatch import invoke_tool

        mcp_name = parse_mcp_checkbox_value(tool_name)
        if mcp_name:
            try:
                from mcp_server import TOOL_REGISTRY

                info = TOOL_REGISTRY.get(mcp_name)
                if not info:
                    return idx, tool_name, f"<pre>MCP tool not found: {mcp_name}</pre>", True
                svc = ""
                if query_analysis and not query_analysis.get("is_general_query"):
                    svc = (query_analysis.get("service_name") or "").strip()
                args = build_mcp_tool_arguments(
                    mcp_name,
                    user_query=input_text,
                    timerange_hours=timerange,
                    service_filter=svc,
                    pagerduty_filters=pagerduty_filters,
                )
                args["_flask_session"] = session
                res = invoke_tool(mcp_name, args, info["function"])
                display = f"MCP:{mcp_name}"
                return idx, display, res, False
            except Exception as e:
                return idx, tool_name, f"<pre>Error executing MCP '{mcp_name}': {e}</pre>", True

        func = TOOLS.get(tool_name, {}).get('function')
        if not func:
            return idx, tool_name, f"<pre>No tool found for {tool_name}</pre>", True
        
        try:
            tool_input = _tool_input_for_request(input_text, tool_name, query_analysis)
            if tool_input and tool_name not in _TOOLS_KEEP_FULL_QUERY and tool_name not in _TOOLS_GLOBAL_SCOPE:
                logging.info("   → %s input: %r", tool_name, tool_input)
            
            if tool_name in _TOOLS_WITH_TIMERANGE:
                res = func(tool_input, timerange)
            elif tool_name == 'Ask_Bedrock':
                res = func(input_text, selected_tools=selected_tools, enable_mcp_access=True)
            elif tool_name == 'Bedrock_Report':
                res = func(tool_input, context_from_other_tools=context_from_other_tools)
            else:
                res = func(tool_input)
            return idx, tool_name, res, False
        except Exception as e:
            return idx, tool_name, f"<pre>Error executing '{tool_name}': {e}</pre>", True
    
    # Separate tools into data tools and synthesis tools
    synthesis_tools = ['Bedrock_Report', 'Ask_Bedrock']
    data_tool_indices = [(idx, tool) for idx, tool in enumerate(selected_tools) if tool not in synthesis_tools]
    synthesis_tool_indices = [(idx, tool) for idx, tool in enumerate(selected_tools) if tool in synthesis_tools]
    
    query_analysis = None
    if input_text.strip():
        query_analysis = analyze_query_with_bedrock(input_text)
        resolved = _monitoring_tool_input_from_analysis(input_text, query_analysis)
        logging.info(
            "🤖 Query analysis: general=%s service=%r resolved=%r confidence=%s",
            query_analysis.get("is_general_query"),
            query_analysis.get("service_name", ""),
            resolved,
            query_analysis.get("confidence", ""),
        )
    
    # Phase 1: Execute data tools in parallel
    context_for_synthesis = {}
    if data_tool_indices:
        logging.info(f"📊 Phase 1: Executing {len(data_tool_indices)} data tool(s) in parallel...")
        with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
            future_to_tool = {
                executor.submit(execute_tool, idx, tool_name, None, query_analysis): (idx, tool_name)
                for idx, tool_name in data_tool_indices
            }
            
            for future in concurrent.futures.as_completed(future_to_tool):
                idx, tool_name, result, is_error = future.result()
                with results_lock:
                    results_dict[idx] = (tool_name, result, is_error)
                    # Store ALL results for synthesis tools (no filtering)
                    if not is_error:
                        context_for_synthesis[tool_name] = result
                        logging.info(f"✅ {tool_name} completed - adding to context")
        logging.info(f"✅ Phase 1 complete: {len(context_for_synthesis)} tool(s) executed")
    
    # Phase 2: Execute synthesis tools with context from data tools
    if synthesis_tool_indices:
        logging.info(f"🧠 Phase 2: Executing {len(synthesis_tool_indices)} synthesis tool(s) with context from {len(context_for_synthesis)} data tool(s)...")
        with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
            future_to_tool = {
                executor.submit(
                    execute_tool,
                    idx,
                    tool_name,
                    context_for_synthesis if context_for_synthesis else None,
                    query_analysis,
                ): (idx, tool_name)
                for idx, tool_name in synthesis_tool_indices
            }
            
            for future in concurrent.futures.as_completed(future_to_tool):
                idx, tool_name, result, is_error = future.result()
                with results_lock:
                    results_dict[idx] = (tool_name, result, is_error)
        logging.info(f"✅ Phase 2 complete")
    
    # Build tabs and results - show ALL tools (no filtering)
    tabs = []
    results = []
    
    for idx in range(len(selected_tools)):
        if idx in results_dict:
            tool_name, res, is_error = results_dict[idx]
            
            # Create tab button
            tab_id = f"tool-tab-{idx}"
            content_id = f"tool-content-{idx}"
            # Set Bedrock_Report as active tab, otherwise first tab
            is_active = "active" if (tool_name == 'Bedrock_Report' or (idx == 0 and 'Bedrock_Report' not in [selected_tools[i] for i in range(len(selected_tools)) if i in results_dict])) else ""
            
            tabs.append(f"""
                <button class='tab-btn {is_active}' onclick='switchTab("{content_id}", this)' data-tab='{content_id}'>
                    {tool_name}
                </button>
            """)
            
            # Create tab content - wrap in container to ensure proper isolation
            display_style = "block" if is_active else "none"
            results.append(f"""
                <div class='tab-content' id='{content_id}' style='display: {display_style}; position: relative; overflow: hidden;'>
                    <div class='tab-content-wrapper'>
                        {res}
                    </div>
                </div>
            """)
    
    exec_time = round(time.time() - start, 2)
    
    # Build tabs container
    tabs_html = f"""
    <div class='tabs-container'>
        <div class='tabs-header'>
            {''.join(tabs)}
        </div>
        <div class='tabs-body'>
            {''.join(results)}
        </div>
    </div>
    """
    final_result = tabs_html
    logging.info(f"✅ Built UI with {len(tabs)} tab(s)")
    
    # Create a descriptive query name for history
    if input_text.strip():
        history_query = input_text
    else:
        # If no input text, use the tool names
        if len(selected_tools) == 1:
            history_query = selected_tools[0]
        else:
            history_query = " + ".join(selected_tools)
    
    add_to_history(history_query, final_result)
    return jsonify({'result': final_result, 'exec_time': exec_time})

@flask_app.route('/api/alerts', methods=['POST'])
def api_alerts():
    data = request.json
    alert_text = data.get('text', '')
    alert_id = len(alerts_db) + 1
    priority = classify_alert(alert_text)
    cause = identify_cause(alert_text)
    alert = {'id': alert_id, 'text': alert_text, 'priority': priority, 'ack': False, 'cause': cause}
    alerts_db.append(alert)
    return jsonify({'status': 'received', 'alert': alert})

@flask_app.route('/api/alerts/ack/<int:alert_id>', methods=['POST'])
def api_ack(alert_id):
    for alert in alerts_db:
        if alert['id'] == alert_id:
            alert['ack'] = True
            return jsonify({'status': 'acknowledged', 'alert': alert})
    return jsonify({'error': 'alert not found'}), 404

@flask_app.route('/api/alerts/status')
def api_alert_status():
    return jsonify(alerts_db)

@flask_app.route('/api/download/docx', methods=['POST'])
def download_docx():
    """Generate and download results as Word document with screenshot image"""
    if not DOCX_AVAILABLE:
        return jsonify({'error': 'Document generation not available. Please install python-docx'}), 503
    
    try:
        data = request.json
        screenshot_image = data.get('screenshot_image', '')
        
        if not screenshot_image:
            return jsonify({'error': 'No screenshot provided'}), 400
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('OneView GOC AI Results', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        time_para = doc.add_paragraph(f'Generated: {timestamp}')
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        time_para.runs[0].font.size = Pt(10)
        time_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # Add spacing
        
        # Decode and insert screenshot image
        try:
            # Remove the data:image/png;base64, prefix
            img_data = screenshot_image.split(',')[1]
            img_bytes = base64.b64decode(img_data)
            
            # Add image to document (full width)
            img_stream = io.BytesIO(img_bytes)
            doc.add_picture(img_stream, width=Inches(6.5))
            
        except Exception as e:
            logging.error(f"Could not add screenshot image: {e}")
            return jsonify({'error': f'Failed to process screenshot: {str(e)}'}), 500
        
        # Save document to BytesIO
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Generate filename
        filename = f"arlo_agenticai_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        logging.error(f"Error generating document: {str(e)}")
        return jsonify({'error': f'Failed to generate document: {str(e)}'}), 500

@flask_app.route('/api/status/monitor')
def api_status_monitor():
    """Endpoint for automatic status monitoring - returns compact status info"""
    try:
        import requests
        from bs4 import BeautifulSoup
        
        url = "https://status.arlo.com"
        resp = requests.get(url, timeout=10)
        
        if resp.status_code != 200:
            return jsonify({'error': f'HTTP {resp.status_code}'})
        
        soup = BeautifulSoup(resp.text, "html.parser")
        text = soup.get_text("\n", strip=True)
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        
        # Extract summary - be more flexible with patterns
        summary = "Status unknown"
        for l in lines:
            l_lower = l.lower()
            if "operational" in l_lower or "all systems operational" in l_lower:
                summary = l
                break
            elif "experiencing issues" in l_lower or "some systems" in l_lower:
                summary = l
                break
            elif "degraded" in l_lower or "partial outage" in l_lower or "major outage" in l_lower:
                summary = l
                break
        
        logging.info(f"📊 Arlo Status Summary: {summary}")
        
        # Extract core services (deduplicate)
        core_services = []
        seen_services = set()
        for i, l in enumerate(lines):
            if l in ["Log In","Notifications","Library","Live Streaming","Video Recording","Arlo Store","Community"]:
                if i+1 < len(lines) and l not in seen_services:
                    status = lines[i+1]
                    # Skip if next line is also a service name (means status wasn't captured)
                    if status not in ["Log In","Notifications","Library","Live Streaming","Video Recording","Arlo Store","Community"]:
                        logging.info(f"✅ Arlo Status: {l} → {status}")
                        core_services.append({"service": l, "status": status})
                        seen_services.add(l)
                    else:
                        logging.warning(f"⚠️ Arlo Status: {l} → status not found (next line is another service: {status})")
        
        # Extract past incidents (last 7 only)
        past_incidents = []
        for i, l in enumerate(lines):
            if any(day in l.lower() for day in ["monday","tuesday","wednesday","thursday","friday","saturday","sunday"]):
                if i+1 < len(lines) and len(past_incidents) < 7:
                    past_incidents.append({"date": l, "detail": lines[i+1]})
        
        return jsonify({
            'summary': summary,
            'services': core_services,
            'incidents': past_incidents,
            'timestamp': time.strftime('%H:%M:%S')
        })
    except Exception as e:
        error_msg = str(e)
        # Simplify proxy errors
        if 'ProxyError' in error_msg or 'Tunnel connection failed' in error_msg or '403 Forbidden' in error_msg:
            error_msg = 'Proxy blocked (check network settings)'
        elif 'Max retries exceeded' in error_msg:
            error_msg = 'Connection timeout (check network)'
        elif 'Connection refused' in error_msg:
            error_msg = 'Service unavailable'
        return jsonify({'error': error_msg})

def _pagerduty_external_status_incidents_url(dashboard_id: str, tab=None):
    """Public UI URL: …/external-status-dashboard/{id}/incidents?tab=… (active=ongoing, resolved, pending)."""
    raw = (os.getenv("PAGERDUTY_SUBDOMAIN") or "arlo").strip()
    raw = raw.replace("https://", "").replace("http://", "").split("/")[0]
    sub = (raw.split(".")[0] if raw else "arlo") or "arlo"
    base = f"https://{sub}.pagerduty.com/external-status-dashboard/{dashboard_id}/incidents"
    if tab in ("active", "resolved", "pending"):
        base += f"?tab={tab}"
    return base


def _samsung_status_dashboard_id():
    """External status board id for Samsung widget (REST: status_dashboard_ids[])."""
    v = os.getenv("SAMSUNG_STATUS_DASHBOARD_ID")
    if v is None:
        return "PRBJIO4"
    s = str(v).strip()
    if not s or s.lower() in ("off", "false", "no", "0", "none", "*"):
        return None
    return s


def _adt_status_dashboard_id():
    """External status board id for ADT status widget (REST: status_dashboard_ids[])."""
    v = os.getenv("ADT_STATUS_DASHBOARD_ID")
    if v is None:
        return "PK1QF1G"
    s = str(v).strip()
    if not s or s.lower() in ("off", "false", "no", "0", "none", "*"):
        return None
    return s


def _partner_status_dashboard_id(env_key: str) -> str | None:
    """External status board id for partner env widgets (CAT, Comcast, etc.)."""
    v = os.getenv(f"{env_key.upper()}_STATUS_DASHBOARD_ID")
    if v is None:
        return None
    s = str(v).strip()
    if not s or s.lower() in ("off", "false", "no", "0", "none", "*"):
        return None
    return s


def _cat_status_dashboard_id():
    return _partner_status_dashboard_id("CAT")


def _comcast_status_dashboard_id():
    return _partner_status_dashboard_id("COMCAST")


def _pagerduty_monitor_payload(dashboard_id):
    """
    Same PagerDuty logic as the status monitor semaphore: get_pagerduty_status_counts /
    build_pagerduty_monitor_api_payload in tools.status_monitor (24h window, total=true,
    retries, SQLite cache). dashboard_id=None = whole account; else external status board.
    """
    try:
        api_token = os.getenv("PAGERDUTY_API_TOKEN")
        if not api_token:
            return {"error": "PagerDuty token not configured"}
        from tools.status_monitor import build_pagerduty_monitor_api_payload

        payload = build_pagerduty_monitor_api_payload(api_token, dashboard_id)
        if payload.get("error"):
            return payload
        if dashboard_id:
            payload["status_dashboard_id"] = dashboard_id
            payload["status_dashboard_url"] = _pagerduty_external_status_incidents_url(dashboard_id)
            payload["status_dashboard_url_active"] = _pagerduty_external_status_incidents_url(
                dashboard_id, tab="active"
            )
            payload["status_dashboard_url_resolved"] = _pagerduty_external_status_incidents_url(
                dashboard_id, tab="resolved"
            )
            payload["status_dashboard_url_pending"] = _pagerduty_external_status_incidents_url(
                dashboard_id, tab="pending"
            )
        return payload
    except Exception as e:
        logging.error(f"Error in PagerDuty monitor payload: {e}")
        error_msg = str(e)
        if "ProxyError" in error_msg or "Tunnel connection failed" in error_msg or "403 Forbidden" in error_msg:
            error_msg = "Proxy blocked (check network settings)"
        elif "Max retries exceeded" in error_msg:
            error_msg = "Connection timeout (check network)"
        elif "Connection refused" in error_msg:
            error_msg = "Service unavailable"
        return {"error": error_msg}


def _jsonify_pagerduty_monitor(data: dict):
    if data.get("error"):
        return jsonify(data)
    return jsonify(data)


@flask_app.route('/api/pagerduty/monitor')
def api_pagerduty_monitor():
    """All-account PagerDuty incidents (no external status board filter)."""
    return _jsonify_pagerduty_monitor(_pagerduty_monitor_payload(None))


@flask_app.route('/api/pagerduty/incidents')
def api_pagerduty_incidents():
    """PagerDuty incidents table for query UI (shift / missing RCA toggles)."""
    from tools.pagerduty_tool import get_pagerduty_incidents
    from tools.pagerduty_team import normalize_pagerduty_shift

    query = (request.args.get("query") or "").strip()
    shift = normalize_pagerduty_shift(request.args.get("shift"))
    team_only = (request.args.get("team_only") or "0").strip().lower() in ("1", "true", "yes", "on")
    missing_rca = (request.args.get("missing_rca") or "0").strip().lower() in ("1", "true", "yes", "on")
    try:
        html = get_pagerduty_incidents(
            query,
            shift=shift,
            team_only=team_only,
            missing_root_cause=missing_rca,
        )
        return jsonify({"html": html})
    except Exception as e:
        logging.error(f"api_pagerduty_incidents: {e}")
        return jsonify({"error": str(e)}), 500


@flask_app.route('/api/pagerduty/samsung-monitor')
def api_pagerduty_samsung_monitor():
    """
    Samsung external status board. Default: scrape public tab HTML (Ongoing / Pending / Resolved).
    Set SAMSUNG_PAGERDUTY_USE_API=1 to use the Incidents REST API instead.
    """
    bid = _samsung_status_dashboard_id()
    if not bid:
        return jsonify(
            {
                "error": "Samsung status disabled",
                "disabled": True,
                "hint": "Set SAMSUNG_STATUS_DASHBOARD_ID (default PRBJIO4) or remove it to use the default.",
            }
        )
    use_api = (os.getenv("SAMSUNG_PAGERDUTY_USE_API") or "").strip().lower() in (
        "1",
        "true",
        "yes",
        "on",
    )
    if use_api:
        return _jsonify_pagerduty_monitor(_pagerduty_monitor_payload(bid))
    try:
        from tools.pagerduty_samsung_scrape import build_samsung_pagerduty_scrape_payload

        pl = build_samsung_pagerduty_scrape_payload(bid)
        pl["status_dashboard_id"] = bid
        pl["status_dashboard_url"] = _pagerduty_external_status_incidents_url(bid)
        pl["status_dashboard_url_active"] = _pagerduty_external_status_incidents_url(bid, tab="active")
        pl["status_dashboard_url_resolved"] = _pagerduty_external_status_incidents_url(
            bid, tab="resolved"
        )
        pl["status_dashboard_url_pending"] = _pagerduty_external_status_incidents_url(
            bid, tab="pending"
        )
        return _jsonify_pagerduty_monitor(pl)
    except Exception as e:
        logging.warning("Samsung PagerDuty scrape failed, falling back to API: %s", e)
        return _jsonify_pagerduty_monitor(_pagerduty_monitor_payload(bid))


@flask_app.route('/api/pagerduty/adt-monitor')
def api_pagerduty_adt_monitor():
    """ADT external status board: scoped to ADT_STATUS_DASHBOARD_ID (default PK1QF1G)."""
    bid = _adt_status_dashboard_id()
    if not bid:
        return jsonify(
            {
                "error": "ADT status disabled",
                "disabled": True,
                "hint": "Set ADT_STATUS_DASHBOARD_ID (default PK1QF1G) or remove it to use the default.",
            }
        )
    return _jsonify_pagerduty_monitor(_pagerduty_monitor_payload(bid))


@flask_app.route('/api/pagerduty/cat-monitor')
def api_pagerduty_cat_monitor():
    """CAT external status board: scoped to CAT_STATUS_DASHBOARD_ID."""
    bid = _cat_status_dashboard_id()
    if not bid:
        return jsonify(
            {
                "error": "CAT status disabled",
                "disabled": True,
                "hint": "Set CAT_STATUS_DASHBOARD_ID to the PagerDuty external status board id.",
            }
        )
    return _jsonify_pagerduty_monitor(_pagerduty_monitor_payload(bid))


@flask_app.route('/api/pagerduty/comcast-monitor')
def api_pagerduty_comcast_monitor():
    """Comcast external status board: scoped to COMCAST_STATUS_DASHBOARD_ID."""
    bid = _comcast_status_dashboard_id()
    if not bid:
        return jsonify(
            {
                "error": "Comcast status disabled",
                "disabled": True,
                "hint": "Set COMCAST_STATUS_DASHBOARD_ID to the PagerDuty external status board id.",
            }
        )
    return _jsonify_pagerduty_monitor(_pagerduty_monitor_payload(bid))


@flask_app.route("/api/slack/send-results", methods=["POST"])
def api_slack_send_results():
    """Send result to Slack (Incoming Webhook): plain text or Block Kit with mrkdwn (no raw HTML)."""
    webhook = (os.getenv("SLACK_WEBHOOK_URL") or "").strip()
    if not webhook:
        return jsonify(
            {"success": False, "error": _SLACK_WEBHOOK_MISSING_MSG}
        ), 503

    payload = request.get_json(silent=True) or {}
    text = payload.get("text")
    if text is None:
        text = ""
    if not isinstance(text, str):
        text = str(text)
    text = text.strip()

    mrkdwn_raw = payload.get("mrkdwn")
    mrkdwn_body = ""
    if isinstance(mrkdwn_raw, str):
        mrkdwn_body = mrkdwn_raw.strip()

    if not text and not mrkdwn_body:
        return jsonify({"success": False, "error": "No text to send"}), 400

    max_len = 38000
    if len(mrkdwn_body) > max_len:
        mrkdwn_body = mrkdwn_body[:max_len] + "\n\n_[... truncated for Slack ...]_"
    if mrkdwn_body and len(text) > 4000:
        text = text[:3997] + "..."

    try:
        if mrkdwn_body:
            fallback = text or (mrkdwn_body[:500] + ("…" if len(mrkdwn_body) > 500 else ""))
            if len(fallback) > 4000:
                fallback = fallback[:3997] + "..."
            chunk_size = 2900
            blocks = [
                {
                    "type": "header",
                    "text": {"type": "plain_text", "text": "OneView GOC AI", "emoji": True},
                },
                {"type": "divider"},
            ]
            i = 0
            while i < len(mrkdwn_body) and len(blocks) < 49:
                blocks.append(
                    {
                        "type": "section",
                        "text": {"type": "mrkdwn", "text": mrkdwn_body[i : i + chunk_size]},
                    }
                )
                i += chunk_size
            if i < len(mrkdwn_body):
                blocks.append(
                    {
                        "type": "section",
                        "text": {"type": "mrkdwn", "text": "_[... additional content omitted (Slack limit) ...]_"},
                    }
                )
            slack_payload = {"text": fallback, "blocks": blocks}
        else:
            main = text
            if len(main) > max_len:
                main = main[:max_len] + "\n\n[... truncated for Slack ...]"
            slack_payload = {"text": main}

        r = post_incoming_webhook(webhook, slack_payload, timeout=(15, 45))
        if r.status_code != 200:
            logging.warning("Slack webhook HTTP %s: %s", r.status_code, (r.text or "")[:300])
            return jsonify(
                {
                    "success": False,
                    "error": f"Slack returned HTTP {r.status_code}",
                }
            ), 502
        return jsonify({"success": True})
    except Exception as e:
        logging.error("Slack webhook: %s", e)
        return jsonify({"success": False, "error": format_slack_connection_error(e)}), 500


def _slack_safe_mrkdwn_snippet(s: str, max_len: int = 400) -> str:
    """Avoid breaking Slack mrkdwn in context lines (* and _ are special)."""
    if not s:
        return ""
    t = str(s).replace("&", "&amp;").replace("*", "·").replace("_", " ").replace("`", "'")
    t = " ".join(t.split())
    return t[:max_len] + ("…" if len(t) > max_len else "")


def _slack_summary_mrkdwn_chunks(summary: str, max_chunk: int = 2900) -> list[str]:
    """Split on paragraphs when possible so mrkdwn (*bold*) is less often cut mid-token."""
    summary = (summary or "").strip()
    if not summary:
        return []
    if len(summary) <= max_chunk:
        return [summary]
    paras = summary.split("\n\n")
    out = []
    buf = ""
    for p in paras:
        p = p.strip()
        if not p:
            continue
        candidate = (buf + "\n\n" + p) if buf else p
        if len(candidate) <= max_chunk:
            buf = candidate
            continue
        if buf:
            out.append(buf)
            buf = ""
        if len(p) <= max_chunk:
            buf = p
        else:
            start = 0
            while start < len(p):
                out.append(p[start : start + max_chunk])
                start += max_chunk
    if buf:
        out.append(buf)
    return out


def _slack_summary_blocks_payload(summary: str, page_title: str, source_url: str) -> dict:
    """Block Kit: header, context, divider, section(s) with mrkdwn for a richer Slack message."""
    chunks = _slack_summary_mrkdwn_chunks(summary, 2900)
    if not chunks:
        chunks = [""]

    fallback = "📊 OneView — AI summary (Bedrock)"
    if page_title:
        fallback += " · " + _slack_safe_mrkdwn_snippet(page_title, 120)

    blocks = [
        {
            "type": "header",
            "text": {"type": "plain_text", "text": "📊 OneView · AI summary", "emoji": True},
        },
    ]

    ctx_lines = []
    if page_title:
        ctx_lines.append(f"📌 {_slack_safe_mrkdwn_snippet(page_title, 350)}")
    if source_url:
        su = source_url.strip().replace("|", "%7C")
        ctx_lines.append(f"<{su}|Open in browser>")
    if ctx_lines:
        blocks.append(
            {
                "type": "context",
                "elements": [{"type": "mrkdwn", "text": "\n".join(ctx_lines)}],
            }
        )

    blocks.append({"type": "divider"})

    for part in chunks:
        part = (part or "").strip()
        if not part:
            continue
        if len(blocks) >= 48:
            blocks.append(
                {
                    "type": "section",
                    "text": {"type": "mrkdwn", "text": "_…message truncated for Slack…_"},
                }
            )
            break
        blocks.append({"type": "section", "text": {"type": "mrkdwn", "text": part}})

    return {"text": fallback[:4000], "blocks": blocks}


@flask_app.route("/api/slack/summarize-and-send", methods=["POST"])
def api_slack_summarize_and_send():
    """
    Plain text from client → AWS Bedrock summary (English, mrkdwn-friendly) → Slack webhook with Block Kit.
    Requires SLACK_WEBHOOK_URL and BEDROCK_API_KEY.
    """
    webhook = (os.getenv("SLACK_WEBHOOK_URL") or "").strip()
    if not webhook:
        return jsonify(
            {"success": False, "error": _SLACK_WEBHOOK_MISSING_MSG}
        ), 503

    data = request.get_json(silent=True) or {}
    raw = data.get("page_text")
    if raw is None:
        raw = ""
    if not isinstance(raw, str):
        raw = str(raw)
    page_text = raw.strip()

    pt = data.get("page_title")
    page_title = pt.strip() if isinstance(pt, str) else ""
    su = data.get("source_url")
    source_url = su.strip() if isinstance(su, str) else ""

    if not page_text:
        return jsonify({"success": False, "error": "page_text is empty"}), 400

    try:
        max_in = int(os.getenv("SLACK_SUMMARY_INPUT_MAX_CHARS", "100000"))
    except (TypeError, ValueError):
        max_in = 100000
    max_in = max(5000, min(max_in, 200000))
    if len(page_text) > max_in:
        page_text = page_text[:max_in] + "\n\n[... content truncated for summary ...]"

    try:
        from tools.bedrock_tool import summarize_page_for_slack

        summary = summarize_page_for_slack(page_text, page_title, source_url)
    except Exception as e:
        logging.exception("Slack summarize Bedrock")
        return jsonify({"success": False, "error": str(e)}), 500

    err_markers = (
        summary.startswith("Error:"),
        summary.startswith("AWS Bedrock Error"),
        "Bedrock is not working" in summary,
    )
    if any(err_markers):
        return jsonify({"success": False, "error": summary[:2000]}), 502

    slack_payload = _slack_summary_blocks_payload(summary, page_title, source_url)

    try:
        r = post_incoming_webhook(webhook, slack_payload, timeout=(15, 90))
        if r.status_code != 200:
            logging.warning("Slack webhook HTTP %s: %s", r.status_code, (r.text or "")[:300])
            return jsonify(
                {"success": False, "error": f"Slack returned HTTP {r.status_code}"},
            ), 502
        return jsonify({"success": True})
    except Exception as e:
        logging.error("Slack summarize webhook post: %s", e)
        return jsonify({"success": False, "error": format_slack_connection_error(e)}), 500


@flask_app.route("/api/extension/chat", methods=["POST"])
def api_extension_chat():
    """
    Browser extension chat: quick (Wiki + Bedrock) or deep (GocBedrock / Bedrock_Report).
    Returns answer in JSON and posts summary to Slack.
    """
    webhook = (os.getenv("SLACK_WEBHOOK_URL") or "").strip()
    if not webhook:
        return jsonify({"success": False, "error": _SLACK_WEBHOOK_MISSING_MSG}), 503

    data = request.get_json(silent=True) or {}
    input_text = (data.get("input") or "").strip()
    if not input_text:
        return jsonify({"success": False, "error": "input is required"}), 400

    mode = str(data.get("mode") or "deep").strip().lower()
    if mode not in ("quick", "deep"):
        return jsonify({"success": False, "error": "mode must be quick or deep"}), 400

    source_url = data.get("source_url")
    source_url = source_url.strip() if isinstance(source_url, str) else ""
    page_title = (
        f"GocView Quick — {input_text[:120]}"
        if mode == "quick"
        else f"GocView Deep — {input_text[:120]}"
    )

    from tools.issues_context import html_to_plain_text
    from tools.bedrock_tool import summarize_page_for_slack

    start = time.time()
    try:
        if mode == "quick":
            html_result = _extension_quick_search(input_text)
        else:
            if not ARLOCHAT_AVAILABLE:
                return jsonify(
                    {"success": False, "error": "GocBedrock (Bedrock_Report) is not available"},
                ), 503
            html_result = ask_arlo(input_text)
    except Exception as e:
        logging.exception("Extension chat mode=%s", mode)
        return jsonify({"success": False, "error": str(e)}), 500

    answer_text = html_to_plain_text(html_result)
    if not answer_text.strip():
        answer_text = input_text

    page_text = answer_text
    try:
        max_in = int(os.getenv("SLACK_SUMMARY_INPUT_MAX_CHARS", "100000"))
    except (TypeError, ValueError):
        max_in = 100000
    max_in = max(5000, min(max_in, 200000))
    if len(page_text) > max_in:
        page_text = page_text[:max_in] + "\n\n[... content truncated for summary ...]"

    try:
        summary = summarize_page_for_slack(page_text, page_title, source_url)
    except Exception as e:
        logging.exception("Extension chat Slack summarize")
        return jsonify({"success": False, "error": str(e)}), 500

    err_markers = (
        summary.startswith("Error:"),
        summary.startswith("AWS Bedrock Error"),
        "Bedrock is not working" in summary,
    )
    if any(err_markers):
        return jsonify({"success": False, "error": summary[:2000]}), 502

    slack_payload = _slack_summary_blocks_payload(summary, page_title, source_url)
    slack_sent = False
    try:
        r = post_incoming_webhook(webhook, slack_payload, timeout=(15, 90))
        if r.status_code != 200:
            logging.warning("Extension Slack HTTP %s: %s", r.status_code, (r.text or "")[:300])
            return jsonify(
                {"success": False, "error": f"Slack returned HTTP {r.status_code}"},
            ), 502
        slack_sent = True
    except Exception as e:
        logging.error("Extension Slack webhook: %s", e)
        return jsonify({"success": False, "error": format_slack_connection_error(e)}), 500

    exec_time = round(time.time() - start, 2)
    return jsonify(
        {
            "success": True,
            "mode": mode,
            "message": "Answer ready; summary sent to Slack",
            "exec_time": exec_time,
            "slack_sent": slack_sent,
            "answer_html": html_result,
            "answer_text": answer_text[:12000],
        }
    )


@flask_app.route("/api/shift/report", methods=["POST"])
def api_shift_report():
    """
    Shift handoff table: PagerDuty, Datadog, Slack, GRM calendar + Jira GRM tickets for a fixed shift window.
    mode: shift1 (Mexico 11:30–20:00) | shift2 (20:00–02:30) | shift3 (02:30–11:30) Mexico time
    """
    data = request.get_json(silent=True) or {}
    mode = str(data.get("mode") or "shift1").strip().lower()
    try:
        from tools.shift_report import _normalize_shift_mode

        mode = _normalize_shift_mode(mode)
    except ValueError as e:
        return jsonify({"success": False, "error": str(e)}), 400

    send_slack = bool(data.get("send_slack", False))
    start = time.time()

    try:
        from tools.shift_report import generate_shift_report
        from tools.issues_context import html_to_plain_text
        from tools.bedrock_tool import summarize_page_for_slack

        report = generate_shift_report(mode)
        html_result = report["html"]
        answer_text = report.get("plain_text") or html_to_plain_text(html_result)
        exec_time = round(time.time() - start, 2)

        slack_sent = False
        if send_slack:
            webhook = (os.getenv("SLACK_WEBHOOK_URL") or "").strip()
            if not webhook:
                return jsonify(
                    {
                        "success": True,
                        "result": html_result,
                        "summary": report.get("summary"),
                        "row_count": report.get("row_count", 0),
                        "csv": report.get("csv"),
                        "mode": mode,
                        "label": report["label"],
                        "window_start": report["window_start"],
                        "window_end": report["window_end"],
                        "timezone": report["timezone"],
                        "sources": report["sources"],
                        "exec_time": exec_time,
                        "slack_sent": False,
                        "slack_error": _SLACK_WEBHOOK_MISSING_MSG,
                    }
                ), 200

            page_title = f"GocView Shift — {report['label']}"
            summary = summarize_page_for_slack(answer_text, page_title, "")
            slack_payload = _slack_summary_blocks_payload(summary, page_title, "")
            r = post_incoming_webhook(webhook, slack_payload, timeout=(15, 90))
            if r.status_code != 200:
                return jsonify(
                    {
                        "success": True,
                        "result": html_result,
                        "summary": report.get("summary"),
                        "row_count": report.get("row_count", 0),
                        "csv": report.get("csv"),
                        "mode": mode,
                        "label": report["label"],
                        "exec_time": exec_time,
                        "slack_sent": False,
                        "slack_error": f"Slack returned HTTP {r.status_code}",
                    }
                ), 200
            slack_sent = True

        return jsonify(
            {
                "success": True,
                "result": html_result,
                "summary": report.get("summary"),
                "row_count": report.get("row_count", 0),
                "csv": report.get("csv"),
                "mode": mode,
                "label": report["label"],
                "window_start": report["window_start"],
                "window_end": report["window_end"],
                "timezone": report["timezone"],
                "sources": report["sources"],
                "exec_time": exec_time,
                "slack_sent": slack_sent,
            }
        )
    except Exception as e:
        logging.exception("Shift report mode=%s", mode)
        return jsonify({"success": False, "error": str(e)}), 500


def _extension_quick_search(input_text: str) -> str:
    """
    Onboarding-style quick lookup: Confluence Wiki + short Bedrock answer (no MCP / deep triage).
    """
    from tools.issues_context import html_to_plain_text

    wiki_html = confluence_search(input_text)
    wiki_plain = html_to_plain_text(wiki_html)
    if len(wiki_plain) > 12000:
        wiki_plain = wiki_plain[:12000] + "\n... (wiki truncated)"

    prompt = f"""You are GocView Quick Assist (onboarding-style documentation helper for Arlo GOC).

User question: "{input_text}"

Confluence / wiki search results (primary source — cite page titles when relevant):
{wiki_plain or "(no wiki pages matched)"}

TASK:
- Give a concise, practical answer in HTML only (<div>...</div>).
- Prefer documentation links and runbook-style guidance from the wiki results.
- Do NOT invent live metrics, incidents, or ticket IDs not present above.
- Keep it short: hero line + up to 5 bullet findings or steps.
- This is NOT a deep incident triage — no recurrence analysis or long tool dumps.
- No markdown code fences.

Return ONLY the HTML."""

    return ask_bedrock(prompt, selected_tools=None, enable_mcp_access=False)


@flask_app.route("/api/slack/screenshot-image/<sid>", methods=["GET"])
def api_slack_screenshot_image(sid):
    """Temporary PNG URL for Slack (Incoming Webhook) to fetch image_url."""
    if not re.match(r"^[a-f0-9]{32}$", sid):
        return jsonify({"error": "invalid id"}), 400
    data = _slack_screenshot_get_png(sid)
    if not data:
        return jsonify({"error": "expired or not found"}), 404
    return Response(
        data,
        mimetype="image/png",
        headers={"Cache-Control": "no-store, max-age=0"},
    )


@flask_app.route("/api/slack/send-screenshot", methods=["POST"])
def api_slack_send_screenshot():
    """
    Screenshots via the same Slack Incoming Webhook as text:
    stores the PNG for a few minutes and posts an attachment with image_url (Slack fetches the URL).

    The app must be reachable from the internet over HTTPS (set PUBLIC_BASE_URL if needed).
    Optional: files.upload only if SLACK_BOT_TOKEN starts with xoxb- and SLACK_CHANNEL_ID is set.
    """
    if "file" not in request.files:
        return jsonify({"success": False, "error": "No file received (field: file)"}), 400
    upload = request.files["file"]
    if not upload or not upload.filename:
        return jsonify({"success": False, "error": "Empty file"}), 400

    data = upload.read()
    max_bytes = 8 * 1024 * 1024
    if len(data) > max_bytes:
        return jsonify({"success": False, "error": "Image too large (max 8 MB)"}), 400

    caption = (request.form.get("caption") or "").strip() or "OneView GOC AI — screenshot"
    caption = caption[:500]

    webhook = (os.getenv("SLACK_WEBHOOK_URL") or "").strip()

    try:
        if webhook:
            sid = _slack_screenshot_store_png(data)
            base = _slack_screenshot_public_base_url()
            if not base:
                return jsonify(
                    {"success": False, "error": "Could not resolve public base URL for this server"}), 500
            image_url = f"{base}/api/slack/screenshot-image/{sid}"
            # Slack only embeds image_url over public HTTPS. With http://localhost the message
            # still delivers if we include the link in text (open PNG in browser).
            if image_url.startswith("https://"):
                payload = {
                    "text": caption,
                    "attachments": [
                        {
                            "image_url": image_url,
                            "fallback": "OneView screenshot",
                        }
                    ],
                }
            else:
                logging.warning(
                    "screenshot Slack: URL not HTTPS (%s) — sending link in text (Slack does not embed HTTP). "
                    "In production set PUBLIC_BASE_URL=https://your-domain",
                    image_url[:160],
                )
                payload = {
                    "text": (
                        f"{caption}\n\n"
                        f"PNG screenshot (temporary ~{_SLACK_SCREENSHOT_TTL_SEC}s): {image_url}"
                    ),
                }
            r = post_incoming_webhook(webhook, payload, timeout=(15, 45))
            if r.status_code != 200:
                logging.warning(
                    "Slack webhook screenshot HTTP %s: %s",
                    r.status_code,
                    (r.text or "")[:400],
                )
                return jsonify(
                    {
                        "success": False,
                        "error": f"Slack returned HTTP {r.status_code}: {(r.text or '')[:200]}",
                    }
                ), 502
            return jsonify({"success": True})

        token = (os.getenv("SLACK_BOT_TOKEN") or "").strip()
        channel = (
            os.getenv("SLACK_CHANNEL_ID") or os.getenv("SLACK_SCREENSHOT_CHANNEL") or ""
        ).strip()
        if token.startswith("xoxb-") and channel:
            files_part = {
                "file": ("oneview-screenshot.png", io.BytesIO(data), "image/png"),
            }
            r = post_slack_api(
                "https://slack.com/api/files.upload",
                headers={"Authorization": f"Bearer {token}"},
                data={
                    "channels": channel,
                    "initial_comment": caption,
                },
                files=files_part,
                timeout=(15, 90),
            )
            try:
                body = r.json()
            except Exception:
                logging.warning("Slack files.upload no-JSON: %s", (r.text or "")[:400])
                return jsonify(
                    {"success": False, "error": f"Slack returned HTTP {r.status_code}"}
                ), 502
            if not body.get("ok"):
                err = body.get("error") or "unknown_error"
                logging.warning("Slack files.upload error: %s", err)
                return jsonify(
                    {"success": False, "error": f"Slack: {err}"}
                ), 502
            return jsonify({"success": True})

        return jsonify(
            {
                "success": False,
                "error": (
                    "Set SLACK_WEBHOOK_URL (recommended for screenshots) or a xoxb- bot + SLACK_CHANNEL_ID. "
                    "In Docker use --env-file .env or env_file in compose; the image does not ship .env."
                ),
            }
        ), 503
    except Exception as e:
        logging.error("Slack screenshot upload: %s", e)
        return jsonify({"success": False, "error": format_slack_connection_error(e)}), 500


@flask_app.route("/api/piranha/probe")
def api_piranha_probe():
    """Check Piranha Okta/ALB session (no secrets returned)."""
    from tools.piranha_employees import piranha_auth_status

    return jsonify(piranha_auth_status(session))


@flask_app.route("/api/piranha/session", methods=["POST", "DELETE"])
def api_piranha_session():
    """Save or clear Piranha browser session cookies (Okta via ALB)."""
    from tools.piranha_employees import piranha_auth_status
    from tools.piranha_session import clear_cookies, save_cookies, validate_session
    from tools.servicenow_session import cookie_header_from_dict, parse_cookie_blob

    if request.method == "DELETE":
        clear_cookies(session)
        return jsonify({"success": True, "auth": piranha_auth_status(session)})

    body = request.get_json(silent=True) or {}
    cookies: dict[str, str] = {}
    raw_header = ""

    if isinstance(body.get("cookies"), dict):
        cookies = {str(k): str(v) for k, v in body["cookies"].items() if v}
        raw_header = cookie_header_from_dict(cookies) if cookies else ""
    else:
        raw = str(body.get("cookie") or body.get("cookies") or "").strip()
        cookies = parse_cookie_blob(raw)
        raw_header = raw

    if not cookies and not raw_header:
        return jsonify(
            {
                "success": False,
                "error": "Missing cookie header. Paste AWSELBAuthSessionCookie-0/1 from piranha.arlo.com.",
            }
        ), 400

    save_cookies(session, cookies, raw_header=raw_header or None)
    ok, err = validate_session(session)
    if not ok:
        clear_cookies(session)
        return jsonify({"success": False, "error": err, "auth": piranha_auth_status(session)}), 401

    return jsonify({"success": True, "auth": piranha_auth_status(session)})


@flask_app.route("/api/piranha/connect/auto/start", methods=["POST"])
def api_piranha_auto_start():
    try:
        from tools.piranha_browser_connect import start_auto_connect

        out = start_auto_connect()
        if out.get("success") and out.get("connect_id"):
            session["piranha_connect_id"] = out["connect_id"]
            session.modified = True
        return jsonify(out), (200 if out.get("success") else 503)
    except Exception as e:
        logging.error("Piranha auto-connect start: %s", e)
        return jsonify({"success": False, "error": str(e)}), 500


@flask_app.route("/api/piranha/connect/auto/status")
def api_piranha_auto_status():
    try:
        from tools.piranha_browser_connect import poll_auto_connect

        connect_id = (
            (request.args.get("connect_id") or "").strip()
            or (session.get("piranha_connect_id") or "").strip()
        )
        if not connect_id:
            return jsonify({"status": "unknown", "error": "No connection in progress."})
        out = poll_auto_connect(connect_id, session)
        if out.get("status") == "connected":
            session.pop("piranha_connect_id", None)
        return jsonify(out)
    except Exception as e:
        logging.error("Piranha auto-connect status: %s", e)
        return jsonify({"status": "error", "error": str(e)})


@flask_app.route("/api/sentinel/certificates")
def api_sentinel_certificates():
    """Sidebar: SSL certificate semaphore from sentinel.arlocloud.com."""
    try:
        from tools.sentinel_certificates import sentinel_certificates_payload

        query = (request.args.get("query") or "").strip()
        force = request.args.get("refresh") in ("1", "true", "yes")
        return jsonify(sentinel_certificates_payload(query, force_refresh=force))
    except Exception as e:
        logging.error("Error in Sentinel certificates: %s", e)
        return jsonify({"success": False, "error": str(e), "summary": {}, "expired": [], "expiring": []})


@flask_app.route("/api/splunk/monitor")
def api_splunk_monitor():
    """Sidebar: P0 Splunk tools — outlier counts per zone (predict band), same SPL as chat tools."""
    try:
        from tools.splunk_tool import splunk_outliers_monitor_payload

        tr = request.args.get("timerange", type=int)
        return jsonify(splunk_outliers_monitor_payload(tr))
    except Exception as e:
        logging.error("Error in Splunk monitor: %s", e)
        return jsonify(
            {
                "success": False,
                "error": str(e),
                "tools": [],
            }
        )


@flask_app.route("/api/servicenow/probe")
def api_servicenow_probe():
    """Check server/env ServiceNow session (no secrets returned)."""
    from tools.servicenow_oauth import auth_status
    from tools.servicenow_session import server_env_auth_available, validate_session

    auth = auth_status(session)
    env = server_env_auth_available()
    ok, err = validate_session(None if env else session)
    return jsonify(
        {
            "auth": auth,
            "server_env_configured": env,
            "rest_ok": ok,
            "error": err or None,
        }
    )


@flask_app.route("/api/servicenow/dashboard")
def api_servicenow_dashboard():
    """Sidebar: ServiceDesk KPIs + chart data from ServiceNow REST API."""
    try:
        from tools.servicenow_dashboard import servicedesk_dashboard_payload

        return jsonify(servicedesk_dashboard_payload(session))
    except Exception as e:
        logging.error("Error in ServiceNow dashboard: %s", e)
        return jsonify({"success": False, "error": str(e), "kpis": {}})


@flask_app.route("/api/servicenow/auth")
def api_servicenow_auth():
    """OAuth connection status for ServiceDesk widget."""
    try:
        from tools.servicenow_oauth import auth_status

        return jsonify(auth_status(session))
    except Exception as e:
        logging.error("Error in ServiceNow auth status: %s", e)
        return jsonify({"configured": False, "connected": False, "error": str(e)})


@flask_app.route("/oauth/snow/login")
def oauth_snow_login():
    """Start ServiceNow OAuth (Okta) login."""
    from urllib.parse import quote

    from tools.servicenow_oauth import build_authorize_url, oauth_redirect_uri

    try:
        return_to = (request.args.get("next") or "/").strip() or "/"
        url = build_authorize_url(
            session,
            redirect_uri=oauth_redirect_uri(request.url_root),
            return_to=return_to,
        )
        return redirect(url)
    except Exception as e:
        logging.error("ServiceNow OAuth login: %s", e)
        return redirect("/?snow=error&msg=" + quote(str(e)[:200]))


@flask_app.route("/oauth/snow/callback")
def oauth_snow_callback():
    """OAuth redirect from ServiceNow after Okta login."""
    from urllib.parse import quote

    from tools.servicenow_oauth import (
        RETURN_KEY,
        STATE_KEY,
        exchange_code_for_tokens,
        oauth_redirect_uri,
    )

    err = (request.args.get("error") or "").strip()
    if err:
        desc = (request.args.get("error_description") or err).strip()
        return redirect("/?snow=error&msg=" + quote(desc[:200]))

    state = (request.args.get("state") or "").strip()
    expected = (session.get(STATE_KEY) or "").strip()
    if not state or state != expected:
        return redirect("/?snow=error&msg=" + quote("Invalid OAuth state — please try again."))

    code = (request.args.get("code") or "").strip()
    if not code:
        return redirect("/?snow=error&msg=" + quote("Missing OAuth code."))

    try:
        exchange_code_for_tokens(
            session,
            code=code,
            redirect_uri=oauth_redirect_uri(request.url_root),
        )
    except Exception as e:
        logging.error("ServiceNow OAuth callback: %s", e)
        return redirect("/?snow=error&msg=" + quote(str(e)[:200]))

    return_to = session.pop(RETURN_KEY, "/") or "/"
    sep = "&" if "?" in return_to else "?"
    return redirect(f"{return_to}{sep}snow=connected")


@flask_app.route("/api/servicenow/connect/auto/start", methods=["POST"])
def api_servicenow_auto_start():
    """Open browser via Playwright; user logs in with Okta; cookies captured automatically."""
    try:
        from tools.servicenow_browser_connect import start_auto_connect

        out = start_auto_connect()
        if out.get("success") and out.get("connect_id"):
            session["snow_connect_id"] = out["connect_id"]
            session.modified = True
        return jsonify(out), (200 if out.get("success") else 503)
    except Exception as e:
        logging.error("ServiceNow auto-connect start: %s", e)
        return jsonify({"success": False, "error": str(e)}), 500


@flask_app.route("/api/servicenow/connect/auto/status")
def api_servicenow_auto_status():
    """Poll Playwright auto-connect progress."""
    try:
        from tools.servicenow_browser_connect import poll_auto_connect

        connect_id = (
            (request.args.get("connect_id") or "").strip()
            or (session.get("snow_connect_id") or "").strip()
        )
        if not connect_id:
            return jsonify({"status": "unknown", "error": "No connection in progress."})
        out = poll_auto_connect(connect_id, session)
        if out.get("status") == "connected":
            session.pop("snow_connect_id", None)
        return jsonify(out)
    except Exception as e:
        logging.error("ServiceNow auto-connect status: %s", e)
        return jsonify({"status": "error", "error": str(e)})


@flask_app.route("/api/servicenow/session", methods=["POST", "DELETE"])
def api_servicenow_session():
    """Save or clear ServiceNow browser session cookie (manual Okta login)."""
    from tools.servicenow_oauth import auth_status
    from tools.servicenow_session import (
        clear_cookies,
        cookie_header_from_dict,
        parse_cookie_blob,
        save_session_auth,
        validate_session,
    )

    if request.method == "DELETE":
        clear_cookies(session)
        from tools.servicenow_oauth import clear_token_bundle

        clear_token_bundle(session)
        return jsonify({"success": True, "auth": auth_status(session)})

    body = request.get_json(silent=True) or {}
    cookies: dict[str, str] = {}
    raw_header = ""
    user_token = str(body.get("user_token") or body.get("g_ck") or body.get("gck") or "").strip()

    if isinstance(body.get("cookies"), dict):
        cookies = {str(k): str(v) for k, v in body["cookies"].items() if v}
        raw_header = cookie_header_from_dict(cookies) if cookies else ""
    else:
        jsessionid = (body.get("jsessionid") or body.get("JSESSIONID") or "").strip()
        glide = (body.get("glide_session_store") or body.get("glide") or "").strip()
        route = (body.get("glide_user_route") or body.get("user_route") or "").strip()
        if jsessionid:
            cookies["JSESSIONID"] = jsessionid
        if glide:
            cookies["glide_session_store"] = glide
        if route:
            cookies["glide_user_route"] = route
        if cookies:
            raw_header = cookie_header_from_dict(cookies)
        else:
            raw = str(body.get("cookie") or body.get("cookies") or "").strip()
            cookies = parse_cookie_blob(raw)
            raw_header = raw

    if not user_token:
        parsed = parse_cookie_blob(raw_header) if raw_header else cookies
        user_token = str(parsed.get("g_ck") or parsed.get("X-UserToken") or "").strip()

    if not cookies and not raw_header:
        return jsonify(
            {
                "success": False,
                "error": "Paste JSESSIONID, glide_session_store, and g_ck token.",
            }
        ), 400

    if not user_token:
        return jsonify(
            {
                "success": False,
                "error": "Missing g_ck token. In ServiceNow: Develop → Web Inspector → Console → window.g_ck",
            }
        ), 400

    save_session_auth(session, cookies, raw_header=raw_header, user_token=user_token)
    ok, err = validate_session(session)
    if not ok:
        clear_cookies(session)
        return jsonify({"success": False, "error": err}), 401

    return jsonify({"success": True, "auth": auth_status(session)})


@flask_app.route("/oauth/snow/logout")
def oauth_snow_logout():
    """Clear ServiceNow OAuth tokens and session cookies."""
    from tools.servicenow_oauth import clear_token_bundle
    from tools.servicenow_session import clear_cookies

    clear_token_bundle(session)
    clear_cookies(session)
    return redirect("/?snow=logged_out")


@flask_app.route('/api/public-ip', methods=['GET'])
def get_public_ip():
    """Public egress IP plus hints to reach this app (SSH, curl, base URL)."""
    try:
        import socket

        import requests

        services = [
            'https://api.ipify.org?format=json',
            'https://ipinfo.io/json',
            'https://ifconfig.me/all.json',
        ]

        ip = None
        service_used = None
        for service in services:
            try:
                response = requests.get(service, timeout=5)
                if response.status_code == 200:
                    data = response.json()
                    ip = data.get('ip') or data.get('ip_addr')
                    if ip:
                        service_used = service
                        break
            except Exception:
                continue

        if not ip:
            return jsonify({'error': 'Unable to fetch public IP'}), 500

        base = request.url_root.rstrip('/')
        try:
            host_name = socket.gethostname()
        except Exception:
            host_name = ''

        connect = {
            'hostname': host_name,
            'app_base_url': base,
            'note': (
                'This is the server public egress IP (useful for firewall rules). '
                'If you access via load balancer or domain, use the browser URL. '
                'SSH works only if port 22 is open and you have user/key access.'
            ),
            'ssh_examples': [
                {
                    'label': 'SSH (replace user and auth method)',
                    'command': f'ssh <user>@{ip}',
                },
                {
                    'label': 'Amazon Linux / AL2023 (typical on EC2)',
                    'command': f'ssh ec2-user@{ip}',
                },
                {
                    'label': 'Ubuntu cloud AMI',
                    'command': f'ssh ubuntu@{ip}',
                },
            ],
            'curl_health': f'curl -sS "{base}/api/health"',
            'docker_hint': (
                '# If the app runs in Docker on this host (adjust container name):\n'
                '# docker ps\n'
                '# docker exec -it <container_name> /bin/sh'
            ),
        }

        return jsonify({'ip': ip, 'service': service_used, 'connect': connect})
    except Exception as e:
        logging.error(f"Error fetching public IP: {e}")
        return jsonify({'error': str(e)}), 500


_TEAM_CAL_SUBCAL_RESOLVE_TTL_S = 1800.0
_team_cal_subcal_resolve_cache: dict = {}


def _team_calendar_base_candidates() -> list:
    """Bases for Team Calendars /events.json. Cloud often needs /wiki/rest/...; some tenants only respond without /wiki/."""
    override = (os.getenv("DEPLOYMENTS_TEAM_CALENDAR_BASE") or "").strip().rstrip("/")
    if override:
        return [override]
    h = (os.getenv("CONFLUENCE_ATLASSIAN_HOST") or "https://arlo.atlassian.net").strip().rstrip("/")
    if not h.startswith("http"):
        h = f"https://{h.lstrip('/')}"
    return [
        f"{h}/wiki/rest/calendar-services/1.0/calendar",
        f"{h}/rest/calendar-services/1.0/calendar",
    ]


def _fetch_subcalendars_from_bases(
    bases, email, token, *, space_key: str | None = None, include_id: str | None = None
) -> tuple[object, str] | tuple[None, list]:
    """(json, base_used) on first HTTP 200 with a non-empty payload, else (None, [status lines])."""
    import requests

    from tools.grm_calendar_browser import deployments_space_key

    auth = (email, token)
    lines = []
    sk = (space_key or deployments_space_key()).strip()

    def _has_payload(data: object) -> bool:
        if not isinstance(data, dict):
            return False
        payload = data.get("payload")
        if isinstance(payload, list) and payload:
            return True
        return bool(_collect_subcalendars_for_match(data))

    for base in bases:
        url = f"{base.rstrip('/')}/subcalendars.json"
        try:
            r = requests.get(url, auth=auth, timeout=22)
            lines.append(f"{base} bare -> {r.status_code}")
            if r.status_code == 200:
                data = r.json()
                if _has_payload(data):
                    return (data, base)
        except Exception as e:
            lines.append(f"{base} -> err {e!s}")

    # Browser widget: spaceCalendars + viewingSpaceKey (+ optional include parent id)
    import time

    for base in bases:
        url = f"{base.rstrip('/')}/subcalendars.json"
        params = {
            "calendarContext": "spaceCalendars",
            "viewingSpaceKey": sk,
            "_": time.time_ns(),
        }
        if include_id:
            params["include"] = include_id
        try:
            r = requests.get(url, auth=auth, params=params, timeout=22)
            lines.append(f"{base} space:{sk} -> {r.status_code}")
            if r.status_code == 200:
                data = r.json()
                if _has_payload(data):
                    return (data, base)
        except Exception as e:
            lines.append(f"{base} space -> err {e!s}")
    return (None, lines)


def _team_calendar_cst_range_iso(
    now_cst, lookback_days: int, fetch_days_ahead: int
) -> tuple[str, str, str, str]:
    """(start_iso, end_iso, start_ymd, end_ymd) for Team Calendar; API expects ISO-8601 bounds."""
    from datetime import timedelta

    s = (now_cst - timedelta(days=lookback_days)).replace(
        hour=0, minute=0, second=0, microsecond=0
    )
    e = (now_cst + timedelta(days=fetch_days_ahead)).replace(
        hour=23, minute=59, second=59, microsecond=0
    )
    return s.isoformat(), e.isoformat(), s.strftime("%Y-%m-%d"), e.strftime("%Y-%m-%d")


def _collect_subcalendars_for_match(obj) -> list:
    """Flatten Team Calendars /subcalendars.json tree to [(id, name), ...]."""
    out: list = []

    def walk(node) -> None:
        if node is None:
            return
        if isinstance(node, list):
            for it in node:
                walk(it)
        elif isinstance(node, dict):
            cid = node.get("id")
            cname = node.get("name") or node.get("summary") or ""
            if cid is not None and str(cname).strip():
                out.append((str(cid), str(cname).strip()))
            for ch in (
                "childSubCalendars",
                "childSubCals",
                "subCalendars",
                "items",
                "children",
            ):
                if ch in node and node[ch] is not None:
                    walk(node[ch])

    walk(obj)
    return out


def _resolve_subcalendar_id_by_name(
    email: str, token: str, name_substr: str
) -> tuple[Optional[str], list[tuple[str, str]]]:
    """
    Query subcalendars.json; match calendars whose name contains name_substr (case-insensitive).
    Returns (first_matching_id, all_matches) for diagnostics when several calendars match.
    """
    if not name_substr or not email or not token:
        return None, []
    name_key = name_substr.strip().lower()
    if not name_key:
        return None, []
    import time

    ent = _team_cal_subcal_resolve_cache.get(name_key)
    if ent and (time.monotonic() - ent[0]) < _TEAM_CAL_SUBCAL_RESOLVE_TTL_S:
        if len(ent) == 3:
            _, cal_id, matches = ent
            return cal_id, list(matches)
        _, cal_id = ent
        return cal_id, []

    bases = _team_calendar_base_candidates()
    data, used_or_lines = _fetch_subcalendars_from_bases(bases, email, token)
    if data is None:
        from tools.grm_calendar_browser import (
            deployments_space_key,
            list_space_subcalendars_for_name_match,
        )

        sk = deployments_space_key()
        data2, base2 = list_space_subcalendars_for_name_match(email, token, sk, bases=bases)
        if data2 is not None:
            data, used_or_lines = data2, base2
    if data is None:
        logging.warning("subcalendars.json: no 200 from any base: %s", used_or_lines)
        return None, []
    logging.info("subcalendars.json OK from base %s", used_or_lines)
    try:
        pairs = _collect_subcalendars_for_match(data)
    except Exception as e:
        logging.warning("subcalendars parse failed: %s", e)
        return None, []

    matches = [(cal_id, cal_name) for cal_id, cal_name in pairs if name_key in cal_name.lower()]
    if not matches:
        logging.warning(
            "No subcalendar matched name containing %r (scanned %s names)",
            name_substr,
            len(pairs),
        )
        return None, []

    cal_id = matches[0][0]
    _team_cal_subcal_resolve_cache[name_key] = (time.monotonic(), cal_id, tuple(matches))
    logging.info(
        "Resolved subCalendarId by name %r -> %s (%s match(es))",
        name_substr,
        cal_id,
        len(matches),
    )
    return cal_id, matches


def _normalize_team_calendar_events(payload, _depth: int = 0) -> list:
    """Team Calendar /events.json may return a list or nested objects; recurse into data/values."""
    if payload is None or _depth > 6:
        return []
    if isinstance(payload, list):
        if not payload:
            return []
        return payload
    if isinstance(payload, dict):
        for key in (
            "events",
            "allEvents",
            "calendarEvents",
            "data",
            "values",
            "results",
            "items",
            "rows",
            "eventList",
            "content",
        ):
            v = payload.get(key)
            if isinstance(v, list):
                return v
            if isinstance(v, dict) and v:
                inner = _normalize_team_calendar_events(v, _depth + 1)
                if inner:
                    return inner
        if any(
            k in payload
            for k in (
                "start",
                "fromDateTime",
                "startDate",
                "startMillis",
                "localStartDate",
                "title",
                "summary",
                "name",
                "what",
            )
        ):
            return [payload]
    return []


def _flatten_team_calendar_event(event: dict) -> dict:
    """Some payloads nest the real event under 'event' / 'payload' / 'item' (Jira/TC wrappers)."""
    if not event:
        return {}
    base = dict(event)
    for wrap in ("event", "item", "payload", "vevent", "value"):
        w = event.get(wrap)
        if isinstance(w, dict):
            for k, v in w.items():
                if v not in (None, "") and (k not in base or base.get(k) in (None, "")):
                    base[k] = v
    return base


def _parse_team_calendar_datetime(val) -> Optional[datetime]:
    """Parse start/end from Team Calendar: ISO string, epoch ms, or nested dict."""
    from datetime import datetime, timezone

    if val is None:
        return None
    if isinstance(val, (int, float)):
        sec = val / 1000.0 if val > 1e12 else float(val)
        try:
            return datetime.fromtimestamp(sec, tz=timezone.utc)
        except (OSError, ValueError, OverflowError):
            return None
    if isinstance(val, dict):
        for k in ("iso", "iso8601", "dateTime", "date", "value", "epochMillis", "millis"):
            if k in val and val[k] not in (None, ""):
                parsed = _parse_team_calendar_datetime(val[k])
                if parsed:
                    return parsed
        return None
    if isinstance(val, str):
        s = val.strip()
        if not s:
            return None
        # All-day events: YYYY-MM-DD only
        if len(s) == 10 and s[4] == "-" and s[7] == "-":
            try:
                from datetime import date as date_cls

                d0 = date_cls.fromisoformat(s)
                return datetime(d0.year, d0.month, d0.day, tzinfo=timezone.utc)
            except ValueError:
                pass
        s = s.replace("Z", "+00:00")
        try:
            dt = datetime.fromisoformat(s)
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            return dt
        except ValueError:
            pass
        try:
            dt = datetime.fromisoformat(s.replace(" ", "T", 1))
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            return dt
        except ValueError:
            pass
        # e.g. "19-Jul-2023" (Confluence all-day in English)
        for fmt in ("%d-%b-%Y", "%d-%B-%Y"):
            try:
                d0 = datetime.strptime(s, fmt).date()
                return datetime(d0.year, d0.month, d0.day, tzinfo=timezone.utc)
            except ValueError:
                continue
        return None
    return None


def _team_calendar_event_start_end(event: dict) -> tuple[Optional[datetime], Optional[datetime]]:
    """Return (start, end) as timezone-aware UTC from various Team Calendar field names."""
    from datetime import timedelta

    event = _flatten_team_calendar_event(event)
    start_dt = None
    end_dt = None
    for sk in (
        "start",
        "startDate",
        "startTime",
        "fromDateTime",
        "from",
        "begin",
        "startMillis",
        "localStartDate",
        "originalStartDate",
    ):
        if sk in event and event[sk] not in (None, ""):
            start_dt = _parse_team_calendar_datetime(event[sk])
            if start_dt:
                break
    if start_dt is None and isinstance(event.get("start"), dict):
        start_dt = _parse_team_calendar_datetime(event["start"])
    for ek in (
        "end",
        "endDate",
        "endTime",
        "toDateTime",
        "to",
        "finish",
        "endMillis",
        "localEndDate",
        "originalEndDate",
    ):
        if ek in event and event[ek] not in (None, ""):
            end_dt = _parse_team_calendar_datetime(event[ek])
            if end_dt:
                break
    if end_dt is None and isinstance(event.get("end"), dict):
        end_dt = _parse_team_calendar_datetime(event["end"])
    if start_dt and end_dt is None:
        end_dt = start_dt + timedelta(hours=2)
    return start_dt, end_dt


def _grm_event_to_deployment(event: dict, cst) -> Optional[dict]:
    """Build one deployment dict from Team Calendar API event (start/end in ISO UTC)."""
    fe = _flatten_team_calendar_event(event)
    title = (
        fe.get("what")
        or fe.get("title")
        or fe.get("summary")
        or fe.get("name")
        or "Untitled Deployment"
    )
    if not isinstance(title, str):
        title = str(title)

    start_dt, end_dt = _team_calendar_event_start_end(fe)
    if start_dt is None:
        return None

    deploy_dt_cst = start_dt.astimezone(cst)
    if end_dt is None:
        from datetime import timedelta

        end_dt = start_dt + timedelta(hours=2)
    end_cst = end_dt.astimezone(cst)
    return {
        "date": deploy_dt_cst.strftime("%b %d, %Y"),
        "service": title,
        "timestamp": deploy_dt_cst.isoformat(),
        "end_timestamp": end_cst.isoformat(),
    }


def _load_team_calendar_events_for_id(
    email: str,
    token: str,
    cst,
    sub_calendar_id: str,
    start_date_iso: str,
    end_date_iso: str,
    start_ymd: str,
    end_ymd: str,
) -> tuple[list, dict]:
    """
    Load GRM events the same way as the Confluence calendar widget: space subcalendars.json
    with include=parent id, then events.json per child sub-calendar. Falls back to a single
    parent subCalendarId request if the browser-style path returns nothing.
    """
    from tools.grm_calendar_browser import (
        deployments_space_key,
        load_calendar_events_browser_style,
    )

    bases = _team_calendar_base_candidates()
    browser_partial: dict = {}
    try:
        raw_events, browser_partial = load_calendar_events_browser_style(
            email,
            token,
            sub_calendar_id,
            start_date_iso,
            end_date_iso,
            space_key=deployments_space_key(),
            normalize_events=_normalize_team_calendar_events,
            bases=bases,
        )
        browser_rows = []
        for event in raw_events:
            try:
                row = _grm_event_to_deployment(event, cst)
                if row:
                    browser_rows.append(row)
            except Exception as e:
                logging.error("Error parsing browser-style event: %s", e)
        if browser_rows:
            browser_partial["raw_events_primary"] = len(raw_events)
            browser_partial["saw_200"] = True
            browser_partial["grm_fetch_mode"] = "browser"
            return browser_rows, browser_partial
    except Exception as e:
        logging.warning("Team Calendar browser-style fetch failed: %s", e)
        browser_partial["browser_error"] = str(e)

    import requests

    auth = (email, token)
    common_iso = {
        "start": start_date_iso,
        "end": end_date_iso,
        "userTimeZoneId": "America/Chicago",
    }
    common_ymd = {
        "start": start_ymd,
        "end": end_ymd,
        "userTimeZoneId": "America/Chicago",
    }
    partial: dict = {
        "primary_http_status": None,
        "alt_http_status": None,
        "saw_200": False,
        "raw_events_primary": 0,
        "raw_events_alt": 0,
        "sample_event_keys": None,
        "calendar_base_used": None,
        "event_fetch_log": [],
        "grm_date_mode_used": "iso",
        "grm_fetch_mode": "legacy_single_id",
    }
    partial.update({k: v for k, v in browser_partial.items() if v is not None})

    def events_from_raw(raw) -> tuple[list, int]:
        evs = _normalize_team_calendar_events(raw)
        rows = []
        for event in evs:
            try:
                row = _grm_event_to_deployment(event, cst)
                if row:
                    rows.append(row)
            except Exception as e:
                logging.error("Error parsing event: %s", e)
        return rows, len(evs)

    def _note_zero_payload(raw, label: str) -> None:
        if not isinstance(partial.get("grm_zero_events_debug"), list):
            partial["grm_zero_events_debug"] = []
        d = [label, type(raw).__name__]
        if isinstance(raw, dict):
            d.append(list(raw.keys())[:20])
        partial["grm_zero_events_debug"].append(d[:3])

    def one_pass(common: dict, mode_label: str) -> list | None:
        for param_name in ("subCalendarId", "calendarId"):
            for base in bases:
                url = f"{base.rstrip('/')}/events.json"
                p = {**common, param_name: sub_calendar_id}
                try:
                    r = requests.get(url, auth=auth, params=p, timeout=25)
                except Exception as e:
                    partial["event_fetch_log"].append(
                        f"{param_name} {mode_label} {base.split('//', 1)[-1][:24]} err {e!s}"
                    )
                    continue
                if param_name == "subCalendarId" and partial["primary_http_status"] is None:
                    partial["primary_http_status"] = r.status_code
                if param_name == "calendarId" and partial["alt_http_status"] is None:
                    partial["alt_http_status"] = r.status_code
                partial["event_fetch_log"].append(
                    f"{param_name} {mode_label} {base.split('//', 1)[-1][:32]} -> {r.status_code}"
                )
                if r.status_code != 200:
                    continue
                partial["saw_200"] = True
                raw = r.json()
                partial["grm_date_mode_used"] = "ymd" if mode_label == "ymd" else "iso"
                partial["calendar_base_used"] = base
                rows, nraw = events_from_raw(raw)
                if param_name == "subCalendarId":
                    partial["raw_events_primary"] = nraw
                else:
                    partial["raw_events_alt"] = nraw
                evs = _normalize_team_calendar_events(raw)
                if evs and isinstance(evs[0], dict):
                    partial["sample_event_keys"] = list(evs[0].keys())[:25]
                if nraw == 0:
                    _note_zero_payload(raw, f"{param_name}/{mode_label}")
                    continue
                if rows:
                    return rows
        return None

    out = one_pass(common_iso, "iso")
    if out is not None:
        return out, partial
    out = one_pass(common_ymd, "ymd")
    if out is not None:
        return out, partial

    if partial.get("saw_200"):
        partial["empty_calendar_200"] = True
    return [], partial


def _deployments_subcalendar_fallback_name() -> str | None:
    """
    If the macro id returns 404, try resolving a subCalendar by this name (default GRM).
    Set DEPLOYMENTS_SUBCALENDAR_ID_FALLBACK_NAME=0 to disable.
    """
    v = (os.getenv("DEPLOYMENTS_SUBCALENDAR_ID_FALLBACK_NAME", "GRM") or "").strip()
    if not v or v.lower() in ("0", "false", "no", "off", "none"):
        return None
    return v


def _deployments_int_env(name: str, default: int, lo: int, hi: int) -> int:
    try:
        v = int((os.getenv(name) or str(default)).strip())
        return max(lo, min(hi, v))
    except (TypeError, ValueError):
        return default


@flask_app.route('/api/deployments/upcoming')
def api_deployments_upcoming():
    """Endpoint for upcoming deployments from Confluence GRM Calendar (Team Calendars API)."""
    try:
        from datetime import datetime, timedelta, timezone
        from zoneinfo import ZoneInfo

        cst = ZoneInfo('America/Chicago')
        use_mock_deployments = os.getenv("DEPLOYMENTS_USE_MOCK_DATA", "").strip().lower() in (
            "1",
            "true",
            "yes",
        )
        
        email = (os.getenv("ATLASSIAN_EMAIL") or "").strip()
        token = (os.getenv("CONFLUENCE_TOKEN") or "").strip()
        # Default = GRM subCalendarId from the calendar macro (ac:parameter name="id"), not the wiki page id.
        sub_calendar_id_env = (os.getenv("DEPLOYMENTS_SUBCALENDAR_ID") or "fb3ba305-784d-4750-a244-db3d87683733").strip()
        name_match = (os.getenv("DEPLOYMENTS_SUBCALENDAR_NAME") or "").strip()
        name_env_resolve_matches: list[tuple[str, str]] = []
        sub_calendar_id = sub_calendar_id_env
        if name_match and email and token:
            maybe_id, name_env_resolve_matches = _resolve_subcalendar_id_by_name(email, token, name_match)
            if maybe_id:
                sub_calendar_id = maybe_id

        if not email or not token:
            return jsonify(
                {
                    "deployments": [],
                    "total": 0,
                    "timestamp": time.strftime("%H:%M:%S"),
                    "source": "no_credentials",
                    "warning": (
                        "Confluence is not configured: set ATLASSIAN_EMAIL and CONFLUENCE_TOKEN "
                        "in .env (or container env) and restart."
                    ),
                }
            )

        today = datetime.now(timezone.utc)

        diag = {
            "sub_calendar_id": sub_calendar_id,
            "sub_calendar_id_env": sub_calendar_id_env,
            "name_match": name_match or None,
            "primary_http_status": None,
            "alt_http_status": None,
            "raw_events_primary": 0,
            "raw_events_alt": 0,
            "calendar_start_date": None,
            "calendar_end_date": None,
            "api_start_param": None,
            "api_end_param": None,
            "sample_event_keys": None,
            "calendar_base_used": None,
            "event_fetch_log": None,
            "saw_200": None,
        }

        deployments = []

        # Date-only range for Team Calendar API (CST). The wiki GRM view shows a multi-week/month
        # horizon; a short end_date used to return almost no events compared to the Confluence page.
        fetch_days_ahead = _deployments_int_env(
            "DEPLOYMENTS_CALENDAR_FETCH_DAYS_AHEAD", 60, 3, 120
        )
        lookback_days = _deployments_int_env("DEPLOYMENTS_CALENDAR_LOOKBACK_DAYS", 2, 0, 30)
        filter_hours_future = _deployments_int_env(
            "DEPLOYMENTS_FILTER_HOURS_FUTURE", 24, 1, 2160
        )
        filter_hours_past = _deployments_int_env("DEPLOYMENTS_FILTER_HOURS_PAST", 0, 0, 168)
        max_rows = _deployments_int_env("DEPLOYMENTS_MAX_ROWS", 50, 5, 150)
        now_cst = datetime.now(cst)
        start_date_iso, end_date_iso, start_ymd, end_ymd = _team_calendar_cst_range_iso(
            now_cst, lookback_days, fetch_days_ahead
        )
        start_date, end_date = start_ymd, end_ymd
        diag["calendar_start_date"] = start_ymd
        diag["calendar_end_date"] = end_ymd
        diag["api_start_param"] = start_date_iso
        diag["api_end_param"] = end_date_iso
        diag["fetch_days_ahead"] = fetch_days_ahead
        diag["filter_hours_future"] = filter_hours_future
        diag["filter_hours_past"] = filter_hours_past
        diag["max_rows"] = max_rows
        if name_env_resolve_matches:
            diag["name_match_candidates"] = [
                {"id": a, "name": b} for a, b in name_env_resolve_matches[:25]
            ]

        # Team Calendars: ISO range + per-base /wiki vs non-wiki + subCalendarId / calendarId; optional name id fallback
        partial: dict = {}
        try:
            deployments, partial = _load_team_calendar_events_for_id(
                email, token, cst, sub_calendar_id, start_date_iso, end_date_iso, start_ymd, end_ymd
            )
            for k in (
                "primary_http_status",
                "alt_http_status",
                "saw_200",
                "raw_events_primary",
                "raw_events_alt",
                "sample_event_keys",
                "calendar_base_used",
                "empty_calendar_200",
                "grm_fetch_mode",
                "deployments_space_key",
                "browser_child_subcalendar_count",
                "browser_raw_events",
                "browser_child_hits",
                "browser_skip_reason",
            ):
                if k in partial and partial[k] is not None:
                    diag[k] = partial[k]
            for k in ("grm_zero_events_debug", "grm_date_mode_used"):
                if k in partial and partial[k] is not None:
                    diag[k] = partial[k]
            if partial.get("event_fetch_log"):
                diag["event_fetch_log"] = partial["event_fetch_log"][:20]
            logging.info(
                "Team Calendar: bases=%s primary=%s alt=%s base_used=%s",
                _team_calendar_base_candidates(),
                partial.get("primary_http_status"),
                partial.get("alt_http_status"),
                partial.get("calendar_base_used"),
            )
        except Exception as e:
            logging.warning("Team Calendar API failed: %s", e)

        fb_name = _deployments_subcalendar_fallback_name()
        # Do not gate on empty_calendar_200: a wrong subCalendarId can still return HTTP 200 with
        # zero events; we must still try resolving "GRM" by name (DEPLOYMENTS_SUBCALENDAR_ID_FALLBACK_NAME).
        if not deployments and fb_name and email and token:
            alt_id, fb_matches = _resolve_subcalendar_id_by_name(email, token, fb_name)
            diag["grm_fallback_name"] = fb_name
            diag["grm_fallback_resolved_id"] = alt_id
            diag["grm_fallback_match_candidates"] = [
                {"id": a, "name": b} for a, b in fb_matches[:25]
            ]
            if not alt_id:
                diag["grm_fallback_skip_reason"] = "subcalendars_unavailable_or_no_name_match"
            elif alt_id == sub_calendar_id:
                diag["grm_fallback_skip_reason"] = (
                    "name_resolved_to_same_id_as_request_events_api_returned_zero_in_range"
                )
            if alt_id and alt_id != sub_calendar_id:
                logging.info(
                    "Calendar: retrying with name %r -> subCalendarId %s (macro id was not accepted)",
                    fb_name,
                    alt_id,
                )
                diag["subcalendar_fallback_name"] = fb_name
                sub_calendar_id = alt_id
                diag["sub_calendar_id"] = sub_calendar_id
                try:
                    deployments, partial2 = _load_team_calendar_events_for_id(
                        email, token, cst, sub_calendar_id, start_date_iso, end_date_iso, start_ymd, end_ymd
                    )
                    for k in (
                        "primary_http_status",
                        "alt_http_status",
                        "saw_200",
                        "raw_events_primary",
                        "raw_events_alt",
                        "sample_event_keys",
                        "calendar_base_used",
                        "empty_calendar_200",
                        "grm_fetch_mode",
                        "deployments_space_key",
                        "browser_child_subcalendar_count",
                        "browser_raw_events",
                        "browser_child_hits",
                        "browser_skip_reason",
                    ):
                        if k in partial2 and partial2[k] is not None:
                            diag[k] = partial2[k]
                    for k in ("grm_zero_events_debug", "grm_date_mode_used"):
                        if k in partial2 and partial2[k] is not None:
                            diag[k] = partial2[k]
                    if partial2.get("event_fetch_log"):
                        diag["event_fetch_log"] = partial2["event_fetch_log"][:20]
                    diag["subcalendar_resolved_by_fallback_name"] = True
                except Exception as e:
                    logging.warning("Team Calendar retry after name match failed: %s", e)

        # Deduplicate (primary + alt may return the same events)
        if deployments:
            seen = set()
            deduped = []
            for d in deployments:
                k = (d.get("service"), d.get("timestamp"))
                if k in seen:
                    continue
                seen.add(k)
                deduped.append(d)
            deployments = deduped

        deployment_source = (
            "calendar_api_browser"
            if deployments and (partial.get("grm_fetch_mode") == "browser" or diag.get("grm_fetch_mode") == "browser")
            else ("calendar_api" if deployments else "empty")
        )
        # Demo rows only when explicitly enabled (never pass as real GRM data)
        if not deployments and use_mock_deployments:
            logging.warning(
                "⚠️ DEPLOYMENTS_USE_MOCK_DATA=1 — generating sample rows (not from GRM calendar)"
            )
            deployment_source = "mock"
            try:
                sample_services = [
                    "HMS Core Services",
                    "Nginx ClientAPI DeviceAPI",
                    "Backend-hmsdevicemanagement",
                    "Advisor Service",
                    "Directory Service",
                    "Backend-hmspubsub",
                    "OAuth Service",
                    "Web Client Release",
                ]
                deployment_times = [-2, -1, 2, 6, 12, 18, 22]
                for idx, hour_offset in enumerate(deployment_times):
                    if idx >= len(sample_services):
                        break
                    deploy_dt = today + timedelta(hours=hour_offset)
                    deploy_dt_cst = deploy_dt.astimezone(cst)
                    end_dt_cst = (today + timedelta(hours=hour_offset + 2)).astimezone(cst)
                    deployments.append(
                        {
                            "date": deploy_dt_cst.strftime("%b %d, %Y"),
                            "service": f"GRM: {sample_services[idx]}",
                            "timestamp": deploy_dt_cst.isoformat(),
                            "end_timestamp": end_dt_cst.isoformat(),
                        }
                    )
                logging.info(f"📂 Generated {len(deployments)} mock deployments")
            except Exception as e:
                logging.error(f"❌ Error generating mock deployments: {e}")
                deployment_source = "empty"
        
        # Overlap with [now - past_window, now + future_window] so the sidebar can mirror the
        # GRM calendar (events beyond 24h were previously never shown).
        past_window = today - timedelta(hours=filter_hours_past)
        next_window = today + timedelta(hours=filter_hours_future)
        filtered_deployments = []

        def _aware(dt):
            if dt.tzinfo is None:
                return dt.replace(tzinfo=timezone.utc)
            return dt

        for deploy in deployments:
            try:
                deploy_start = _aware(
                    datetime.fromisoformat(deploy["timestamp"].replace("Z", "+00:00"))
                )
                if deploy.get("end_timestamp"):
                    deploy_end = _aware(
                        datetime.fromisoformat(deploy["end_timestamp"].replace("Z", "+00:00"))
                    )
                else:
                    deploy_end = deploy_start + timedelta(hours=2)
                    deploy["end_timestamp"] = deploy_end.isoformat()
                # Overlap with filter window
                if deploy_end < past_window or deploy_start > next_window:
                    continue
                deploy["is_past"] = deploy_end < today
                filtered_deployments.append(deploy)
            except Exception as e:
                logging.error(f"Error filtering deployment: {e}")

        deployments = filtered_deployments
        deployments.sort(key=lambda x: x.get("timestamp", ""))
        upcoming = deployments[:max_rows]

        logging.info(
            f"✅ Deployments ({deployment_source}): {len(upcoming)} row(s) in next {filter_hours_future}h window"
        )

        grm_wiki = (
            "https://arlo.atlassian.net/wiki/spaces/RM/pages/153256867/GRM+Calendar"
        )
        payload = {
            "deployments": upcoming,
            "total": len(deployments),
            "timestamp": time.strftime("%H:%M:%S"),
            "source": deployment_source,
            "grm": {
                "confluence_wiki_url": grm_wiki,
                "filter_hours_future": filter_hours_future,
                "filter_hours_past": filter_hours_past,
                "fetch_end_date": end_date,
                "sub_calendar_id": sub_calendar_id,
            },
            "diagnostics": {
                "sub_calendar_id": sub_calendar_id,
                "sub_calendar_id_env": diag.get("sub_calendar_id_env"),
                "name_match": diag.get("name_match"),
                "api_start_param": diag.get("api_start_param"),
                "api_end_param": diag.get("api_end_param"),
                "calendar_start_date": diag.get("calendar_start_date"),
                "calendar_end_date": diag.get("calendar_end_date"),
                "fetch_days_ahead": diag.get("fetch_days_ahead"),
                "filter_hours_future": diag.get("filter_hours_future"),
                "filter_hours_past": diag.get("filter_hours_past"),
                "max_rows": diag.get("max_rows"),
                "primary_http_status": diag.get("primary_http_status"),
                "alt_http_status": diag.get("alt_http_status"),
                "raw_events_primary": diag.get("raw_events_primary"),
                "raw_events_alt": diag.get("raw_events_alt"),
                "saw_200": diag.get("saw_200"),
                "calendar_base_used": diag.get("calendar_base_used"),
                "event_fetch_log": diag.get("event_fetch_log"),
                "subcalendar_resolved_by_fallback_name": diag.get("subcalendar_resolved_by_fallback_name"),
                "rows_after_filter": len(upcoming),
                "sample_event_keys": diag.get("sample_event_keys"),
                "name_match_candidates": diag.get("name_match_candidates"),
                "grm_fallback_name": diag.get("grm_fallback_name"),
                "grm_fallback_resolved_id": diag.get("grm_fallback_resolved_id"),
                "grm_fallback_match_candidates": diag.get("grm_fallback_match_candidates"),
                "grm_fallback_skip_reason": diag.get("grm_fallback_skip_reason"),
                "grm_fetch_mode": diag.get("grm_fetch_mode"),
                "deployments_space_key": diag.get("deployments_space_key"),
                "browser_child_subcalendar_count": diag.get("browser_child_subcalendar_count"),
                "browser_raw_events": diag.get("browser_raw_events"),
                "browser_child_hits": diag.get("browser_child_hits"),
            },
        }
        if deployment_source == "empty":
            ps = diag.get("primary_http_status")
            alt_s = diag.get("alt_http_status")
            raw_total = (diag.get("raw_events_primary") or 0) + (diag.get("raw_events_alt") or 0)
            ok_200 = bool(diag.get("saw_200")) or ps == 200 or alt_s == 200
            if ps in (401, 403) or alt_s in (401, 403):
                payload["warning"] = (
                    "Confluence API returned 401/403: the email/API token pair is invalid or the token "
                    "lacks Confluence access. Create an API token at id.atlassian.com and ensure "
                    "ATLASSIAN_EMAIL matches the Atlassian account."
                )
            elif (ps == 404 or alt_s == 404) and not ok_200:
                payload["warning"] = (
                    f"Calendar API returned 404 (id may not be the Team Calendars subCalendarId). "
                    f"Current: {sub_calendar_id}. The app tries /wiki/rest/ and /rest/; set "
                    "DEPLOYMENTS_SUBCALENDAR_NAME=GRM or set DEPLOYMENTS_SUBCALENDAR_ID from subcalendars.json. "
                    "Set DEPLOYMENTS_SUBCALENDAR_ID_FALLBACK_NAME=0 to skip auto name resolution."
                )
            elif ok_200 and raw_total > 0:
                payload["warning"] = (
                    "Calendar returned events but none could be parsed into deployment rows. "
                    "Confluence may use other field names; check server logs and diagnostics.sample_event_keys."
                )
            elif ok_200 and raw_total == 0:
                sr = diag.get("grm_fallback_skip_reason")
                mode = diag.get("grm_fetch_mode") or partial.get("grm_fetch_mode")
                extra = ""
                if sr == "name_resolved_to_same_id_as_request_events_api_returned_zero_in_range":
                    extra = (
                        " Name lookup resolves to this same id but events.json still returns 0 rows for "
                        "the requested window. If several calendars match, set DEPLOYMENTS_SUBCALENDAR_NAME "
                        "to a more specific substring and inspect diagnostics.grm_fallback_match_candidates."
                    )
                elif sr == "subcalendars_unavailable_or_no_name_match":
                    extra = (
                        " Could not match a sub-calendar by fallback name (check CONFLUENCE_ATLASSIAN_HOST "
                        "and token scope). Set DEPLOYMENTS_SUBCALENDAR_ID from the wiki Team Calendar macro."
                    )
                elif mode == "browser":
                    extra = (
                        " Browser-style fetch (spaceCalendars + child sub-calendars) ran but returned 0 events "
                        f"for DEPLOYMENTS_CONFLUENCE_SPACE_KEY={diag.get('deployments_space_key', 'RM')}."
                    )
                payload["warning"] = (
                    "API returned OK but 0 events in the selected range. Set DEPLOYMENTS_SUBCALENDAR_NAME "
                    "to a substring of the GRM sub-calendar (from subcalendars.json) if the id differs "
                    f"from the wiki page. Current subCalendarId: {sub_calendar_id}."
                ) + extra
            elif ps is not None or alt_s is not None:
                payload["warning"] = (
                    f"Team Calendar API did not return 200 (primary HTTP {ps}, alt HTTP {alt_s}). "
                    "See server logs."
                )
            else:
                payload["warning"] = (
                    "Could not reach Team Calendar API (exception before HTTP response). See server logs. "
                    "Optional: DEPLOYMENTS_USE_MOCK_DATA=1 for demo rows only."
                )
            payload["diagnostics"] = {**payload["diagnostics"], **diag}
        return jsonify(payload)
        
    except Exception as e:
        logging.error(f"Error fetching deployments: {e}")
        error_msg = str(e)
        # Simplify proxy errors
        if 'ProxyError' in error_msg or 'Tunnel connection failed' in error_msg or '403 Forbidden' in error_msg:
            error_msg = 'Proxy blocked (check network settings)'
        elif 'Max retries exceeded' in error_msg:
            error_msg = 'Connection timeout (check network)'
        elif 'Connection refused' in error_msg:
            error_msg = 'Service unavailable'
        return jsonify({'error': error_msg})


# ============================================
# MCP SERVER ENDPOINTS
# ============================================

@flask_app.route('/mcp/sse', methods=['GET', 'POST'])
async def mcp_sse_endpoint():
    """
    MCP Server SSE endpoint
    Exposes OneView GOC AI tools as MCP server for consumption by Claude Desktop, Cursor, etc.
    """
    try:
        from mcp_server import get_mcp_server
        from mcp.server.sse import sse_server
        from starlette.requests import Request as StarletteRequest
        from starlette.responses import StreamingResponse
        
        # Get the MCP server instance
        server = get_mcp_server()
        
        # Convert Flask request to Starlette request format
        # This is needed because MCP SDK uses Starlette
        from werkzeug.datastructures import Headers
        
        # Create a simple adapter
        if request.method == 'GET':
            # For SSE connections
            async def event_stream():
                async with sse_server() as streams:
                    send, receive = streams
                    
                    # Handle the SSE connection
                    async for message in server.handle_sse(receive, send):
                        yield f"data: {json.dumps(message)}\n\n"
            
            return flask_app.response_class(
                event_stream(),
                mimetype='text/event-stream',
                headers={
                    'Cache-Control': 'no-cache',
                    'X-Accel-Buffering': 'no'
                }
            )
        else:
            # Handle POST messages
            data = request.get_json()
            # Process the message through MCP server
            result = await server.handle_request(data)
            return jsonify(result)
            
    except Exception as e:
        logging.error(f"❌ MCP Server endpoint error: {e}")
        return jsonify({'error': str(e)}), 500


@flask_app.route('/mcp/info')
def mcp_info():
    """
    MCP Server information endpoint
    Returns available tools and server metadata
    """
    from mcp_server import TOOL_REGISTRY
    
    return jsonify({
        'name': 'oneview-goc-ai',
        'version': '3.0.0',
        'description': 'OneView GOC AI - Unified monitoring and operations platform',
        'protocol': 'mcp',
        'transport': 'sse',
        'endpoint': '/mcp/sse',
        'tools': [
            {
                'name': name,
                'description': info['description']
            }
            for name, info in TOOL_REGISTRY.items()
        ],
        'total_tools': len(TOOL_REGISTRY)
    })


@flask_app.route('/admin/sql', methods=['GET'])
def sql_console():
    """SQL Console for querying the metrics database"""
    return render_template('sql_console.html')


@flask_app.route('/admin/sql/query', methods=['POST'])
def execute_sql_query():
    """Execute a SQL query against the metrics database (SELECT only)"""
    try:
        data = request.get_json()
        query = data.get('query', '').strip()
        
        if not query:
            return jsonify({'success': False, 'error': 'No query provided'}), 400
        
        # Security: Only allow SELECT queries (read-only)
        query_upper = query.upper().strip()
        if not query_upper.startswith('SELECT'):
            return jsonify({
                'success': False, 
                'error': 'Only SELECT queries are allowed. Queries must start with SELECT.'
            }), 403
        
        # Block dangerous keywords even in SELECT
        dangerous_keywords = ['DROP', 'DELETE', 'UPDATE', 'INSERT', 'ALTER', 'CREATE', 'TRUNCATE', 'REPLACE']
        for keyword in dangerous_keywords:
            if keyword in query_upper:
                return jsonify({
                    'success': False,
                    'error': f'Keyword "{keyword}" is not allowed in queries.'
                }), 403
        
        # Execute query
        import sqlite3
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row  # Enable column name access
        cursor = conn.cursor()
        
        start_time = time.time()
        cursor.execute(query)
        results = cursor.fetchall()
        execution_time = (time.time() - start_time) * 1000  # Convert to ms
        
        # Convert to list of dicts
        columns = [description[0] for description in cursor.description] if cursor.description else []
        rows = [dict(row) for row in results]
        
        conn.close()
        
        return jsonify({
            'success': True,
            'columns': columns,
            'rows': rows,
            'row_count': len(rows),
            'execution_time_ms': round(execution_time, 2)
        })
        
    except sqlite3.Error as e:
        return jsonify({'success': False, 'error': f'SQL Error: {str(e)}'}), 400
    except Exception as e:
        return jsonify({'success': False, 'error': f'Error: {str(e)}'}), 500


if __name__ == '__main__':
    # Default port 8080 (override with env PORT)
    port = int(os.getenv('PORT', 8080))
    
    # Enable threading for concurrent requests
    # This allows multiple users to use the tool simultaneously
    flask_app.run(
        host='0.0.0.0', 
        port=port,
        threaded=True,  # Enable multi-threading for concurrent requests
        debug=False     # Disable debug mode for better performance
    )
